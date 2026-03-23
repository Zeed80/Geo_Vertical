from __future__ import annotations

from collections.abc import Iterable
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.full_report_models import (
    FullReportDraftState,
    FullReportRenderModel,
    OfficialReportContext,
    ReleaseManifest,
    RenderSection,
    RenderTable,
    ValidationIssue,
)
from core.normatives import get_straightness_tolerance, get_vertical_tolerance
from core.report_schema import ContractorInfo, CustomerInfo, EquipmentEntry, FullReportData, Specialist
from core.services.report_templates import ReportDataAssembler, ReportTemplateManager


class FullReportBuilder:
    """Единый builder полного отчета: render-model -> preview/docx."""

    SECTION_TITLES = {
        "title": "1. Реквизиты и титул",
        "participants": "2. Объект и участники",
        "normatives": "3. Нормативная и документальная часть",
        "measurements": "4. Измерения и протоколы",
        "results": "5. Расчеты и выводы",
        "appendices": "6. Приложения и выпуск",
    }

    STRUCTURE_TYPE_LABELS = {
        "tower": "Башня",
        "mast": "Мачта",
        "odn": "ОДН",
    }

    def __init__(self, template_manager: ReportTemplateManager | None = None):
        self.template_manager = template_manager or ReportTemplateManager()

    def build_from_template(
        self,
        template_name: str,
        processed_data,
        raw_data,
        output_path: str | Path,
    ) -> Path:
        draft = self.template_manager.build_draft_from_template(
            template_name,
            processed_data or {},
            raw_data,
        )
        draft.release_manifest = ReleaseManifest(
            output_path=str(output_path),
            exported_at=datetime.now(),
            draft_hash=draft.draft_hash(),
            template_name=template_name,
            included_files=[item.relative_path for item in draft.attachments_manifest if item.include_in_release],
        )
        model = self.assemble_render_model(draft, processed_data, raw_data)
        return self.render_docx(model, output_path)

    def build_docx(self, report: FullReportData, output_path: str | Path) -> Path:
        draft = FullReportDraftState(
            form_data=report,
            official_context=ReportDataAssembler(
                report.calculation_results or {},
                None,
                report.reference_profiles,
            ).build_official_context(report),
        )
        model = self.assemble_render_model(draft, report.calculation_results or {}, None)
        return self.render_docx(model, output_path)

    def assemble_render_model(
        self,
        draft: FullReportDraftState,
        processed_data: dict[str, Any] | None = None,
        raw_data: Any = None,
        import_context: dict[str, Any] | None = None,
        diagnostics: dict[str, Any] | None = None,
    ) -> FullReportRenderModel:
        processed_data = processed_data or {}
        render_draft = draft.copy()
        assembler = ReportDataAssembler(
            processed_data,
            raw_data,
            render_draft.form_data.reference_profiles,
            import_context=import_context,
            import_diagnostics=diagnostics,
        )
        render_draft.form_data = assembler.fill_measurement_sections(render_draft.form_data)
        render_draft.official_context = assembler.build_official_context(
            render_draft.form_data,
            render_draft.official_context,
        )
        issues = self.validate_draft(render_draft, processed_data, raw_data, import_context, diagnostics)
        render_draft.validation_state = issues

        sections = [
            self._build_title_section(render_draft, issues),
            self._build_participants_section(render_draft, issues),
            self._build_normatives_section(render_draft, issues),
            self._build_measurements_section(render_draft, processed_data, import_context, diagnostics, issues),
            self._build_results_section(render_draft, processed_data, issues),
            self._build_appendices_section(render_draft, issues),
        ]

        return FullReportRenderModel(
            document_title="Полный технический отчет",
            subtitle=render_draft.form_data.metadata.project_name or "Без названия объекта",
            report_number=render_draft.form_data.metadata.report_number or "",
            sections=sections,
            warnings=[item.message for item in issues if item.severity == "warning"],
            errors=[item.message for item in issues if item.severity == "error"],
            generated_at=datetime.now(),
        )

    def validate_draft(
        self,
        draft: FullReportDraftState,
        processed_data: dict[str, Any] | None = None,
        raw_data: Any = None,
        import_context: dict[str, Any] | None = None,
        diagnostics: dict[str, Any] | None = None,
    ) -> list[ValidationIssue]:
        processed_data = processed_data or {}
        import_context = import_context or processed_data.get("import_context") or {}
        diagnostics = diagnostics or processed_data.get("import_diagnostics") or {}
        data = draft.form_data
        official = draft.official_context
        issues: list[ValidationIssue] = []

        def add(severity: str, section: str, message: str, field_path: str = "", code: str = ""):
            issues.append(
                ValidationIssue(
                    section=section,
                    severity=severity,
                    message=message,
                    field_path=field_path,
                    code=code,
                )
            )

        if not data.metadata.project_name:
            add("error", "title", "Не заполнено наименование объекта.", "metadata.project_name", "required")
        if not data.metadata.report_number:
            add("error", "title", "Не заполнен номер полного отчета.", "metadata.report_number", "required")
        if official.structure_type not in self.STRUCTURE_TYPE_LABELS:
            add("error", "title", "Укажите корректный тип опоры: башня, мачта или ОДН.", "official.structure_type", "invalid_structure_type")
        if not official.performer:
            add("error", "participants", "Не заполнен исполнитель обследования.", "official.performer", "required")
        if not data.customer.full_name:
            add("error", "participants", "Не заполнен заказчик.", "customer.full_name", "required")
        if not data.contractor.full_name:
            add("error", "participants", "Не заполнена организация-исполнитель.", "contractor.full_name", "required")
        if not official.instrument:
            add("warning", "participants", "Не указан прибор обследования.", "official.instrument", "missing_instrument")
        if not official.weather:
            add("warning", "participants", "Не заполнены погодные условия.", "official.weather", "missing_weather")
        if not official.wind:
            add("warning", "participants", "Не заполнены сведения о ветре.", "official.wind", "missing_wind")
        if data.title_object:
            if data.metadata.inventory_number and data.title_object.inventory_number and data.metadata.inventory_number != data.title_object.inventory_number:
                add("warning", "participants", "Инвентарный номер в карточке объекта отличается от титульного блока.", "title_object.inventory_number", "inventory_mismatch")
            if data.metadata.location and data.title_object.location and data.metadata.location != data.title_object.location:
                add("warning", "participants", "Местоположение в карточке объекта отличается от титульного блока.", "title_object.location", "location_mismatch")

        station_assembler = ReportDataAssembler(
            processed_data,
            raw_data,
            data.reference_profiles,
            import_context=import_context,
            import_diagnostics=diagnostics,
        )
        derived_stations = station_assembler.collect_station_entries()
        if derived_stations:
            actual_names = [item.name.strip() for item in derived_stations if item.name.strip()]
            current_names = [item.name.strip() for item in official.stations if item.name.strip()]
            if current_names and current_names != actual_names:
                add("warning", "participants", "Список стоянок в полном отчёте отличается от standing-point точек проекта.", "official.stations", "station_points_mismatch")
            elif not current_names:
                add("warning", "participants", "Стоянки не заполнены вручную, будут использованы standing-point точки проекта.", "official.stations", "station_points_autofill")
            if official.base_station and official.base_station not in actual_names:
                add("warning", "title", "БС/базовая станция не совпадает с именами standing-point точек проекта.", "official.base_station", "base_station_mismatch")

        vertical = self._build_vertical_summary(draft.form_data, official)
        if not vertical["rows"]:
            add("warning", "measurements", "Нет таблицы отклонений от вертикали для официального протокола.", "vertical_deviation_table", "missing_vertical_protocol")
        if vertical["failed"] > 0 and not official.decision_comment.strip():
            add("error", "results", "Есть превышения по вертикальности, но не заполнено решение/комментарий.", "official.decision_comment", "decision_required")

        confidence = diagnostics.get("confidence", import_context.get("confidence"))
        try:
            confidence_value = float(confidence) if confidence is not None else None
        except (TypeError, ValueError):
            confidence_value = None
        if confidence_value is not None and confidence_value < 0.75:
            add("warning", "measurements", f"Низкая уверенность импорта ({confidence_value:.2f}).", "import.confidence", "low_import_confidence")

        if not draft.attachments_manifest:
            add("warning", "appendices", "Не сформирован манифест вложений.", "attachments_manifest", "attachments_missing")
        return issues

    def render_preview(
        self,
        model: FullReportRenderModel,
        mode: str = "document",
        current_section: str | None = None,
    ) -> str:
        sections = list(model.sections)
        if mode == "section" and current_section:
            sections = [section for section in sections if section.key == current_section]
        elif mode == "errors":
            sections = [section for section in sections if section.errors or section.warnings]

        parts = [
            "<html><head><meta charset='utf-8'><style>",
            "body{font-family:'Segoe UI';background:#ece7de;margin:0;padding:24px;color:#231f1b;}",
            ".page{background:#fff;max-width:980px;margin:0 auto;padding:32px 40px;box-shadow:0 10px 25px rgba(0,0,0,.08);}",
            "h1{font-size:24px;margin:0 0 6px 0;} h2{font-size:18px;margin:28px 0 8px 0;} h3{font-size:15px;margin:18px 0 8px 0;}",
            "p{line-height:1.45;margin:8px 0;} table{width:100%;border-collapse:collapse;margin:10px 0 18px 0;font-size:13px;}",
            "th,td{border:1px solid #b9b0a3;padding:6px 8px;vertical-align:top;} th{background:#f1ece4;}",
            ".meta{color:#6c6357;font-size:12px;margin-bottom:16px;} .error{background:#ffe6e1;border-left:4px solid #c2482c;padding:8px 10px;margin:8px 0;}",
            ".warning{background:#fff1d8;border-left:4px solid #d29100;padding:8px 10px;margin:8px 0;}",
            ".status{font-size:12px;color:#6f655a;text-transform:uppercase;letter-spacing:.06em;}",
            ".prov{font-size:12px;color:#6f655a;margin:6px 0 10px 0;}",
            "</style></head><body><div class='page'>",
            f"<h1>{escape(model.document_title)}</h1>",
            f"<div class='meta'>{escape(model.subtitle)}"
            + (f" | № {escape(model.report_number)}" if model.report_number else "")
            + f" | Сформировано {model.generated_at:%d.%m.%Y %H:%M}</div>",
        ]
        for message in model.errors:
            parts.append(f"<div class='error'>{escape(message)}</div>")
        for message in model.warnings:
            parts.append(f"<div class='warning'>{escape(message)}</div>")
        for section in sections:
            parts.extend(self._render_html_section(section))
        parts.append("</div></body></html>")
        return "".join(parts)

    def render_docx(self, model: FullReportRenderModel, output_path: str | Path) -> Path:
        doc = Document()
        self._configure_styles(doc)
        self._add_header_footer(doc, model)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.space_before = Pt(48)
        title.space_after = Pt(6)
        run = title.add_run(model.document_title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.space_after = Pt(6)
        sub_run = subtitle.add_run(model.subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(12)
        sub_run.font.name = "Times New Roman"

        if model.report_number:
            num_para = doc.add_paragraph()
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = num_para.add_run(f"Отчёт № {model.report_number}")
            nr.font.size = Pt(14)
            nr.bold = True
            nr.font.name = "Times New Roman"

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.space_after = Pt(24)
        dr = date_para.add_run(f"Дата формирования: {model.generated_at:%d.%m.%Y}")
        dr.font.size = Pt(10)
        dr.font.name = "Times New Roman"

        self._add_toc(doc)

        for message in model.errors:
            p = doc.add_paragraph()
            r = p.add_run(f"ОШИБКА: {message}")
            r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            r.bold = True
        for message in model.warnings:
            p = doc.add_paragraph()
            r = p.add_run(f"Предупреждение: {message}")
            r.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

        for section in model.sections:
            doc.add_paragraph(section.title, style="Heading 1")
            for line in section.provenance_summary:
                prov_p = doc.add_paragraph(line)
                prov_p.runs[0].font.size = Pt(8) if prov_p.runs else None
                prov_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88) if prov_p.runs else None
            for message in section.errors:
                p = doc.add_paragraph()
                r = p.add_run(f"Ошибка: {message}")
                r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            for message in section.warnings:
                p = doc.add_paragraph()
                r = p.add_run(f"Предупреждение: {message}")
                r.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
            for paragraph in section.paragraphs:
                para = doc.add_paragraph(paragraph)
                para.paragraph_format.first_line_indent = Cm(1.25)
            for table in section.tables:
                if table.title:
                    doc.add_paragraph(table.title, style="Heading 2")
                self._add_table(doc, table)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _configure_styles(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                h = doc.styles[style_name]
                h.font.name = "Times New Roman"
                h.font.bold = True
                h.font.size = Pt(16 - (level - 1) * 2)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)

        section = doc.sections[0]
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    def _add_header_footer(self, doc: Document, model: FullReportRenderModel) -> None:
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f"Отчёт № {model.report_number}" if model.report_number else "")
        run.font.size = Pt(8)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_number(fp)

    @staticmethod
    def _add_page_number(paragraph) -> None:
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    @staticmethod
    def _add_toc(doc: Document) -> None:
        toc_heading = doc.add_paragraph("СОДЕРЖАНИЕ", style="Heading 1")
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr)

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_separate)

        run2 = p.add_run("(Обновите оглавление в Word: ПКМ → Обновить поле)")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2.italic = True

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run2._r.append(fld_end)

        doc.add_page_break()

    def _render_html_section(self, section: RenderSection) -> list[str]:
        parts = [f"<h2>{escape(section.title)}</h2>", f"<div class='status'>{escape(self._status_label(section.status))}</div>"]
        for line in section.provenance_summary:
            parts.append(f"<div class='prov'>{escape(line)}</div>")
        for message in section.errors:
            parts.append(f"<div class='error'>{escape(message)}</div>")
        for message in section.warnings:
            parts.append(f"<div class='warning'>{escape(message)}</div>")
        for paragraph in section.paragraphs:
            parts.append(f"<p>{escape(paragraph)}</p>")
        for table in section.tables:
            if table.title:
                parts.append(f"<h3>{escape(table.title)}</h3>")
            parts.append("<table><thead><tr>")
            for header in table.headers:
                parts.append(f"<th>{escape(header)}</th>")
            parts.append("</tr></thead><tbody>")
            for row in table.rows:
                parts.append("<tr>")
                for cell in row:
                    parts.append(f"<td>{escape(cell)}</td>")
                parts.append("</tr>")
            parts.append("</tbody></table>")
        return parts

    def _add_table(self, doc: Document, table: RenderTable) -> None:
        if not table.rows:
            doc.add_paragraph("Нет данных для отображения.")
            return
        grid = doc.add_table(rows=1, cols=len(table.headers), style="Table Grid")
        grid.autofit = True
        for index, header in enumerate(table.headers):
            cell = grid.rows[0].cells[index]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
        for row in table.rows:
            cells = grid.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
                for paragraph in cells[index].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"

    def _build_title_section(self, draft: FullReportDraftState, issues: list[ValidationIssue]) -> RenderSection:
        data = draft.form_data
        official = draft.official_context
        rows_main = [
            ["Номер отчета", data.metadata.report_number],
            ["Наименование объекта", data.metadata.project_name],
            ["Инвентарный номер", data.metadata.inventory_number],
            ["Местоположение", data.metadata.location],
            ["Дата обследования", official.survey_date.strftime("%d.%m.%Y")],
            ["Тип опоры", self.STRUCTURE_TYPE_LABELS.get(official.structure_type, official.structure_type)],
        ]
        rows_official = [
            ["БС / базовая станция", official.base_station],
            ["Проект / шифр", official.project_code],
            ["Населенный пункт", official.locality],
            ["Причина измерений", official.measurement_reason],
        ]
        return RenderSection(
            key="title",
            title=self.SECTION_TITLES["title"],
            status=self._resolve_section_status("title", issues, rows_main),
            paragraphs=["Полный отчет сформирован как основной официальный документ проекта."],
            tables=[
                RenderTable("Титульные реквизиты", ["Параметр", "Значение"], rows_main),
                RenderTable("Официальный контекст", ["Параметр", "Значение"], rows_official),
            ],
            errors=self._messages_for_section("title", issues, "error"),
            warnings=self._messages_for_section("title", issues, "warning"),
            required_fields=["Наименование объекта", "Номер отчета", "Тип опоры", "Дата обследования"],
            provenance_summary=self._provenance_summary(draft, "title"),
        )

    def _build_participants_section(self, draft: FullReportDraftState, issues: list[ValidationIssue]) -> RenderSection:
        data = draft.form_data
        official = draft.official_context
        tables = []
        if data.title_object:
            tables.append(
                RenderTable(
                    "Карточка объекта",
                    ["Параметр", "Значение"],
                    [
                        ["Наименование", data.title_object.name],
                        ["Инвентарный номер", data.title_object.inventory_number],
                        ["Эксплуатирующая организация", data.title_object.operator],
                        ["Местоположение", data.title_object.location],
                        ["Год", str(data.title_object.year)],
                    ],
                )
            )
        tables.extend(
            [
                self._make_customer_table(data.customer),
                self._make_contractor_table(data.contractor),
                self._make_specialists_table(data.specialists, official),
                self._make_equipment_table(data.equipment),
            ]
        )
        if official.stations:
            tables.append(
                RenderTable(
                    "Стоянки и расстояния",
                    ["Стоянка", "Расстояние, м", "Примечание"],
                    [[item.name, self._fmt(item.distance_m), item.note] for item in official.stations],
                )
            )
        return RenderSection(
            key="participants",
            title=self.SECTION_TITLES["participants"],
            status=self._resolve_section_status("participants", issues, tables),
            paragraphs=[
                f"Исполнитель обследования: {official.performer or 'не указан'}.",
                f"Проверяющий: {official.reviewer or 'не указан'}.",
                f"Прибор: {official.instrument or 'не указан'}.",
                f"Погодные условия: {official.weather or 'не указаны'}, ветер: {official.wind or 'не указан'}.",
            ],
            tables=tables,
            errors=self._messages_for_section("participants", issues, "error"),
            warnings=self._messages_for_section("participants", issues, "warning"),
            required_fields=["Заказчик", "Организация-исполнитель", "Исполнитель обследования"],
            provenance_summary=self._provenance_summary(draft, "participants"),
        )

    def _build_normatives_section(self, draft: FullReportDraftState, issues: list[ValidationIssue]) -> RenderSection:
        data = draft.form_data
        tables: list[RenderTable] = []
        geodesic = data.geodesic_results or {}
        calculation = data.calculation_results or {}
        summary_rows = [
            ["РљРѕР»РёС‡РµСЃС‚РІРѕ СѓСЂРѕРІРЅРµР№", self._fmt(geodesic.get("total_levels"), 0)],
            ["РњР°РєСЃ. РѕС‚РєР»РѕРЅРµРЅРёРµ, РјРј", self._fmt(geodesic.get("max_deviation_mm"))],
            ["РЎСЂРµРґРЅРµРµ РѕС‚РєР»РѕРЅРµРЅРёРµ, РјРј", self._fmt(geodesic.get("mean_deviation_mm"))],
            ["РњР°РєСЃ. СЃС‚СЂРµР»Р° РїСЂРѕРіРёР±Р°, РјРј", self._fmt(calculation.get("max_straightness_mm"))],
        ]
        summary_rows = [row for row in summary_rows if row[1] not in (None, "")]
        if summary_rows:
            tables.append(RenderTable("РЎРІРѕРґРєР° РїРѕР»СѓС‡РµРЅРЅС‹С… Рё СЂР°СЃС‡РёС‚Р°РЅРЅС‹С… РґР°РЅРЅС‹С…", ["РџР°СЂР°РјРµС‚СЂ", "Р—РЅР°С‡РµРЅРёРµ"], summary_rows))

        level_rows = [
            [
                self._fmt(item.get("index"), 0),
                self._fmt(item.get("height_m")),
                self._fmt(item.get("deviation_mm")),
                self._fmt(item.get("points"), 0),
            ]
            for item in geodesic.get("levels", [])
        ]
        if level_rows:
            tables.append(RenderTable("Р”РµС‚Р°Р»РёР·Р°С†РёСЏ РіРµРѕРґРµР·РёС‡РµСЃРєРёС… СѓСЂРѕРІРЅРµР№", ["РЎРµРєС†РёСЏ", "Р’С‹СЃРѕС‚Р°, Рј", "РћС‚РєР»РѕРЅРµРЅРёРµ, РјРј", "РўРѕС‡РµРє"], level_rows))
        if data.documents:
            tables.append(RenderTable("Исходные документы", ["Документ", "Идентификатор", "Комментарий"], [[item.title, item.identifier, item.comments or ""] for item in data.documents]))
        if data.documents_review:
            tables.append(RenderTable("Анализ документации", ["Документ", "Идентификатор", "Краткий вывод", "Заключение"], [[item.title, item.identifier or "", item.summary or "", item.conclusion or ""] for item in data.documents_review]))
        if data.structural_elements:
            tables.append(RenderTable("Элементы конструкции", ["Секция", "Элемент", "Материал", "Параметры", "Примечание"], [[item.section, item.element, item.material, item.parameters, item.notes or ""] for item in data.structural_elements]))
        paragraphs: list[str] = []
        if data.normative_list:
            paragraphs.append("Нормативная база:")
            paragraphs.extend([f"• {item}" for item in data.normative_list])
        if data.structure:
            paragraphs.extend(
                [
                    data.structure.purpose,
                    data.structure.planning_decisions,
                    data.structure.structural_scheme,
                    data.structure.foundations,
                    data.structure.metal_structure,
                    data.structure.geology or "",
                    data.structure.lattice_notes or "",
                ]
            )
        return RenderSection(
            key="normatives",
            title=self.SECTION_TITLES["normatives"],
            status=self._resolve_section_status("normatives", issues, tables + paragraphs),
            paragraphs=[item for item in paragraphs if item],
            tables=tables,
            errors=self._messages_for_section("normatives", issues, "error"),
            warnings=self._messages_for_section("normatives", issues, "warning"),
            provenance_summary=self._provenance_summary(draft, "normatives"),
        )

    def _build_measurements_section(
        self,
        draft: FullReportDraftState,
        processed_data: dict[str, Any],
        import_context: dict[str, Any] | None,
        diagnostics: dict[str, Any] | None,
        issues: list[ValidationIssue],
    ) -> RenderSection:
        data = draft.form_data
        official = draft.official_context
        import_context = import_context or processed_data.get("import_context") or {}
        diagnostics = diagnostics or processed_data.get("import_diagnostics") or {}
        vertical = self._build_vertical_summary(data, official)
        straightness = self._build_straightness_summary(data)

        tables: list[RenderTable] = []
        geodesic_summary = data.geodesic_results or {}
        calculation_summary = data.calculation_results or {}
        summary_rows: list[list[str]] = []
        if geodesic_summary.get("total_levels") not in (None, ""):
            summary_rows.append(["Количество геодезических уровней", self._fmt(geodesic_summary.get("total_levels"), 0)])
        if geodesic_summary.get("max_deviation_mm") not in (None, ""):
            summary_rows.append(["Максимальное отклонение от вертикали, мм", self._fmt(geodesic_summary.get("max_deviation_mm"))])
        if geodesic_summary.get("mean_deviation_mm") not in (None, ""):
            summary_rows.append(["Среднее отклонение от вертикали, мм", self._fmt(geodesic_summary.get("mean_deviation_mm"))])
        if calculation_summary.get("max_straightness_mm") not in (None, ""):
            summary_rows.append(["Максимальная стрела прогиба, мм", self._fmt(calculation_summary.get("max_straightness_mm"))])
        if summary_rows:
            tables.append(RenderTable("Сводка полученных и рассчитанных данных", ["Параметр", "Значение"], summary_rows))
        level_rows = [
            [
                self._fmt(level.get("index"), 0),
                self._fmt(level.get("height_m")),
                self._fmt(level.get("deviation_mm")),
                self._fmt(level.get("points"), 0),
            ]
            for level in geodesic_summary.get("levels", [])
        ]
        if level_rows:
            tables.append(
                RenderTable(
                    "Детализация геодезических уровней",
                    ["Уровень", "Высота, м", "Отклонение, мм", "Точек"],
                    level_rows,
                )
            )
        if data.measurements:
            tables.append(RenderTable("Сводка измерений", ["Метод", "Норматив", "Результат"], [[item.method, item.standard, item.result] for item in data.measurements]))
        if data.angle_measurements:
            tables.append(
                RenderTable(
                    "Журнал угловых измерений",
                    ["№", "Секция", "Высота, м", "Пояс", "KL", "KR", "Разность", "βизм", "Центр", "Δβ", "Δ, мм"],
                    [
                        [
                            self._fmt(item.index, 0),
                            item.section,
                            self._fmt(item.height_m),
                            item.belt,
                            self._fmt(item.kl_arcsec),
                            self._fmt(item.kr_arcsec),
                            self._fmt(item.diff_arcsec),
                            self._fmt(item.beta_measured),
                            self._fmt(item.center_value),
                            self._fmt(item.delta_beta),
                            self._fmt(item.delta_mm),
                        ]
                        for item in data.angle_measurements
                    ],
                )
            )
        if vertical["rows"]:
            tables.append(RenderTable("Протокол вертикальности", ["Секция", "Высота, м", "Отклонение, мм", "Допуск, мм", "Статус"], vertical["rows"]))
        if straightness["rows"]:
            tables.append(RenderTable("Протокол прямолинейности", ["Пояс", "Высота, м", "Отклонение, мм", "Допуск, мм", "Статус"], straightness["rows"]))
        import_rows = self._build_import_quality_rows(import_context, diagnostics, processed_data)
        if import_rows:
            tables.append(RenderTable("Качество исходных данных и импорта", ["Параметр", "Значение"], import_rows))

        paragraphs = []
        if vertical["conclusion"]:
            paragraphs.append(vertical["conclusion"])
        if straightness["conclusion"]:
            paragraphs.append(straightness["conclusion"])
        return RenderSection(
            key="measurements",
            title=self.SECTION_TITLES["measurements"],
            status=self._resolve_section_status("measurements", issues, tables),
            paragraphs=paragraphs,
            tables=tables,
            errors=self._messages_for_section("measurements", issues, "error"),
            warnings=self._messages_for_section("measurements", issues, "warning"),
            required_fields=["Таблица вертикальности", "Официальный протокол измерений"],
            provenance_summary=self._provenance_summary(draft, "measurements"),
        )

    def _build_results_section(
        self,
        draft: FullReportDraftState,
        processed_data: dict[str, Any],
        issues: list[ValidationIssue],
    ) -> RenderSection:
        data = draft.form_data
        official = draft.official_context
        vertical = self._build_vertical_summary(data, official)
        tables: list[RenderTable] = []
        if data.technical_state:
            tables.append(RenderTable("Оценка технического состояния", ["Конструкция", "Классификация", "Комментарии"], [[item.structure, item.classification, item.comments or ""] for item in data.technical_state]))
        if data.resource_calculation:
            calc = data.resource_calculation
            tables.append(
                RenderTable(
                    "Расчет остаточного ресурса",
                    ["Параметр", "Значение"],
                    [
                        ["Срок службы, лет", self._fmt(calc.service_life_years)],
                        ["Постоянная износа", self._fmt(calc.wear_constant, 4)],
                        ["Полный ресурс, лет", self._fmt(calc.total_service_life_years)],
                        ["Остаточный ресурс, лет", self._fmt(calc.residual_resource_years)],
                        ["Epsilon", self._fmt(calc.epsilon, 4)],
                        ["Lambda", self._fmt(calc.lambda_value, 4)],
                    ],
                )
            )
        paragraphs = self._build_auto_conclusions(draft, processed_data)
        paragraphs.extend([f"{item.label}: {item.text}" for item in data.conclusions])
        if vertical["failed"] > 0 and official.decision_comment:
            paragraphs.append(f"Решение по превышениям: {official.decision_comment}")
        if data.recommendations:
            paragraphs.append("Рекомендации:")
            paragraphs.extend([f"• {item.text}" for item in data.recommendations])
        return RenderSection(
            key="results",
            title=self.SECTION_TITLES["results"],
            status=self._resolve_section_status("results", issues, tables + paragraphs),
            paragraphs=paragraphs,
            tables=tables,
            errors=self._messages_for_section("results", issues, "error"),
            warnings=self._messages_for_section("results", issues, "warning"),
            required_fields=["Заключение по вертикальности", "Итоговые выводы"],
            provenance_summary=self._provenance_summary(draft, "results"),
        )

    def _build_appendices_section(self, draft: FullReportDraftState, issues: list[ValidationIssue]) -> RenderSection:
        data = draft.form_data
        tables: list[RenderTable] = []
        if data.appendices:
            tables.append(RenderTable("Приложения", ["Название", "Описание", "Файлы"], [[item.title, item.description or "", "; ".join(item.files)] for item in data.appendices]))
        if data.annexes:
            tables.append(RenderTable("Перечень приложений A-M", ["Код", "Название", "Описание", "Страницы"], [[item.code, item.title, item.description or "", ", ".join(str(page) for page in item.pages)] for item in data.annexes]))
        if draft.attachments_manifest:
            tables.append(RenderTable("Манифест вложений", ["Название", "Относительный путь", "Выпуск"], [[item.title, item.relative_path, "Да" if item.include_in_release else "Нет"] for item in draft.attachments_manifest]))
        paragraphs = []
        if draft.release_manifest:
            manifest = draft.release_manifest
            paragraphs.extend(
                [
                    f"Последний выпуск: {manifest.output_path}",
                    f"Дата выпуска: {manifest.exported_at:%d.%m.%Y %H:%M}" if manifest.exported_at else "Дата выпуска: не зафиксирована",
                    f"Шаблон: {manifest.template_name or 'не указан'}",
                    f"Хэш версии черновика: {manifest.draft_hash}",
                ]
            )
        return RenderSection(
            key="appendices",
            title=self.SECTION_TITLES["appendices"],
            status=self._resolve_section_status("appendices", issues, tables + paragraphs),
            paragraphs=paragraphs,
            tables=tables,
            errors=self._messages_for_section("appendices", issues, "error"),
            warnings=self._messages_for_section("appendices", issues, "warning"),
            provenance_summary=self._provenance_summary(draft, "appendices"),
        )

    def _build_vertical_summary(self, data: FullReportData, official: OfficialReportContext) -> dict[str, Any]:
        rows: list[list[str]] = []
        failed = 0
        max_excess = 0.0
        structure_type = official.structure_type or "tower"
        for item in data.vertical_deviation_table:
            deviation_mm = item.deviation_current_mm
            if deviation_mm in (None, ""):
                deviation_mm = item.deviation_previous_mm
            if deviation_mm in (None, ""):
                continue
            tolerance_mm = get_vertical_tolerance(float(item.height_m or 0.0), structure_type) * 1000.0
            deviation_value = float(deviation_mm)
            ok = abs(deviation_value) <= tolerance_mm + 1e-9
            if not ok:
                failed += 1
                max_excess = max(max_excess, abs(deviation_value) - tolerance_mm)
            rows.append([self._fmt(item.section_number, 0), self._fmt(item.height_m), self._fmt(deviation_value), self._fmt(tolerance_mm), "Соответствует" if ok else "Превышение"])
        if not rows:
            return {"rows": [], "failed": 0, "conclusion": ""}
        if failed == 0:
            conclusion = f"Вертикальность соответствует нормативному допуску для типа опоры «{self.STRUCTURE_TYPE_LABELS.get(structure_type, structure_type)}»."
        else:
            conclusion = f"Выявлены превышения по вертикальности: {failed}. Максимальное превышение над допуском составляет {self._fmt(max_excess)} мм."
        return {"rows": rows, "failed": failed, "conclusion": conclusion}

    def _build_straightness_summary(self, data: FullReportData) -> dict[str, Any]:
        rows: list[list[str]] = []
        failed = 0
        for item in data.straightness_records:
            tolerance_mm = float(item.tolerance_mm or 0.0)
            if tolerance_mm <= 0 and item.height_m:
                tolerance_mm = get_straightness_tolerance(float(item.height_m)) * 1000.0
            deviation = float(item.deviation_mm or 0.0)
            ok = abs(deviation) <= tolerance_mm + 1e-9
            if not ok:
                failed += 1
            rows.append([self._fmt(item.belt_number, 0), self._fmt(item.height_m), self._fmt(deviation), self._fmt(tolerance_mm), "Соответствует" if ok else "Превышение"])
        if not rows:
            return {"rows": [], "failed": 0, "conclusion": ""}
        return {
            "rows": rows,
            "failed": failed,
            "conclusion": "Отклонения по прямолинейности находятся в пределах допуска." if failed == 0 else f"Выявлены превышения по прямолинейности: {failed}.",
        }

    def _build_import_quality_rows(self, import_context: dict[str, Any], diagnostics: dict[str, Any], processed_data: dict[str, Any]) -> list[list[str]]:
        transformation = processed_data.get("transformation_audit") or {}
        items = [
            ("Формат", import_context.get("source_format") or diagnostics.get("source_format")),
            ("Стратегия парсинга", import_context.get("parser_strategy") or diagnostics.get("parser_strategy")),
            ("Уверенность", self._fmt(import_context.get("confidence", diagnostics.get("confidence")), 2)),
            ("Сырых записей", self._fmt(diagnostics.get("raw_records"), 0)),
            ("Принято точек", self._fmt(diagnostics.get("accepted_points"), 0)),
            ("Отброшено точек", self._fmt(diagnostics.get("discarded_points"), 0)),
            ("RMSE совмещения", self._fmt(transformation.get("rmse"), 4)),
        ]
        rows = [[label, str(value)] for label, value in items if value not in (None, "")]
        warnings = import_context.get("warnings") or diagnostics.get("warnings") or []
        if warnings:
            rows.append(["Предупреждения", " | ".join(str(item) for item in warnings)])
        return rows

    def _build_auto_conclusions(self, draft: FullReportDraftState, processed_data: dict[str, Any]) -> list[str]:
        data = draft.form_data
        official = draft.official_context
        vertical = self._build_vertical_summary(data, official)
        straightness = self._build_straightness_summary(data)
        result = []
        if vertical["conclusion"]:
            result.append(vertical["conclusion"])
        if straightness["conclusion"]:
            result.append(straightness["conclusion"])
        if data.residual_resource:
            status = "обеспечивает безопасную эксплуатацию" if data.residual_resource.satisfies_requirements else "требует дополнительной оценки"
            result.append(f"Остаточный ресурс: {status}, прогнозируемый срок {self._fmt(data.residual_resource.residual_years, 0)} лет.")
        if processed_data.get("import_diagnostics") or processed_data.get("import_context"):
            result.append("В отчет включен блок качества исходных данных и импорта.")
        return result

    def _make_customer_table(self, customer: CustomerInfo) -> RenderTable:
        return RenderTable("Заказчик", ["Параметр", "Значение"], [["Полное наименование", customer.full_name], ["Руководитель", customer.director], ["Юридический адрес", customer.legal_address], ["Фактический адрес", customer.actual_address], ["Телефон", customer.phone], ["E-mail", customer.email]])

    def _make_contractor_table(self, contractor: ContractorInfo) -> RenderTable:
        return RenderTable("Организация-исполнитель", ["Параметр", "Значение"], [["Полное наименование", contractor.full_name], ["Руководитель", contractor.director], ["Юридический адрес", contractor.legal_address], ["Почтовый адрес", contractor.postal_address], ["Телефон", contractor.phone], ["E-mail", contractor.email], ["Аттестация", contractor.accreditation_certificate], ["СРО", contractor.sro_certificate]])

    def _make_specialists_table(self, specialists: Iterable[Specialist], official: OfficialReportContext) -> RenderTable:
        rows = []
        for index, specialist in enumerate(specialists, start=1):
            rows.append([str(index), specialist.full_name, "\n".join(f"{k}: {v}" for k, v in specialist.certifications.items()), "\n".join(f"{k}: {v:%d.%m.%Y}" for k, v in specialist.expires_at.items())])
        if not rows and official.performer:
            rows.append(["1", official.performer, "", ""])
        return RenderTable("Специалисты", ["№", "ФИО", "Аттестации", "Срок действия"], rows)

    def _make_equipment_table(self, equipment: Iterable[EquipmentEntry]) -> RenderTable:
        return RenderTable("Приборы и оборудование", ["№", "Наименование", "Зав. №", "Свидетельство", "Действительно до"], [[str(index), item.name, item.serial_number, item.certificate, item.valid_until.strftime("%d.%m.%Y")] for index, item in enumerate(equipment, start=1)])

    def _messages_for_section(self, section_key: str, issues: list[ValidationIssue], severity: str) -> list[str]:
        return [item.message for item in issues if item.section == section_key and item.severity == severity]

    def _resolve_section_status(self, section_key: str, issues: list[ValidationIssue], content: Any) -> str:
        has_content = False
        if isinstance(content, list):
            has_content = any(bool(item) for item in content)
        else:
            has_content = bool(content)
        if any(item.section == section_key and item.severity == "error" for item in issues):
            return "errors"
        if any(item.section == section_key and item.severity == "warning" for item in issues):
            return "partial"
        if has_content:
            return "ready"
        return "not_started"

    def _status_label(self, status: str) -> str:
        return {
            "errors": "Есть ошибки",
            "partial": "Заполнен частично",
            "ready": "Готов",
            "not_started": "Не начат",
        }.get(status, status)

    def _provenance_summary(self, draft: FullReportDraftState, section_key: str) -> list[str]:
        matches = [
            f"{field}: {source}"
            for field, source in sorted(draft.field_provenance.items())
            if field.startswith(section_key + ".") or field.startswith(section_key + ":")
        ]
        if matches:
            return [f"Источник данных: {', '.join(matches[:4])}"]
        return []

    @staticmethod
    def _fmt(value: Any, precision: int = 2) -> str:
        if value in (None, ""):
            return ""
        try:
            if isinstance(value, float):
                formatted = f"{value:.{precision}f}"
                return formatted.rstrip("0").rstrip(".")
            if isinstance(value, int):
                return str(value)
            numeric = float(value)
            formatted = f"{numeric:.{precision}f}"
            return formatted.rstrip("0").rstrip(".")
        except (TypeError, ValueError):
            return str(value)


__all__ = ["FullReportBuilder"]
