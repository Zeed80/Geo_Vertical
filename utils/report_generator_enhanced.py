"""
Расширенный генератор отчетов с профессиональным оформлением
"""

import logging
import math
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

import numpy as np
import pandas as pd

logger = logging.getLogger(__name__)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, 
                                Spacer, PageBreak, Image, Frame, PageTemplate)
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import tempfile


class EnhancedReportGenerator:
    """Расширенный генератор отчетов для геодезического контроля"""
    
    def __init__(self):
        self.timestamp = datetime.now()
        self.page_width, self.page_height = A4
        self._register_fonts()
    
    @staticmethod
    def _aggregate_angular_measurements_by_sections(angular_measurements: Dict[str, list], 
                                                    height_tolerance: float = 0.3) -> list:
        """
        Агрегирует данные угловых измерений по секциям для получения отклонений по X, Y и суммарного.
        
        Args:
            angular_measurements: Словарь с ключами 'x' и 'y', содержащий списки данных угловых измерений
            height_tolerance: Допуск по высоте для группировки точек в одну секцию (метры)
            
        Returns:
            Список словарей с данными по секциям:
            [
                {
                    'section_num': int,
                    'height': float,
                    'deviation_x': float,  # мм
                    'deviation_y': float,  # мм
                    'total_deviation': float  # мм (sqrt(dev_x^2 + dev_y^2))
                },
                ...
            ]
        """
        if not angular_measurements:
            return []
        
        rows_x = angular_measurements.get('x', [])
        rows_y = angular_measurements.get('y', [])
        
        if not rows_x and not rows_y:
            return []
        
        # Группируем данные по высотам (секциям) для каждой оси
        def group_by_height(rows, tolerance):
            """Группирует строки по высоте с заданной толерантностью"""
            if not rows:
                return {}
            
            groups = {}
            for row in rows:
                height = row.get('height')
                if height is None:
                    continue
                
                delta_mm = row.get('delta_mm')
                if delta_mm is None:
                    continue
                
                # Находим ближайшую группу по высоте
                matched_key = None
                for key_height in groups.keys():
                    if abs(height - key_height) <= tolerance:
                        matched_key = key_height
                        break
                
                if matched_key is not None:
                    groups[matched_key].append(delta_mm)
                else:
                    groups[height] = [delta_mm]
            
            # Вычисляем среднее отклонение для каждой группы
            result = {}
            for height, deviations in groups.items():
                valid_deviations = [d for d in deviations if d is not None]
                if valid_deviations:
                    result[height] = np.mean(valid_deviations)
                else:
                    result[height] = 0.0
            
            return result
        
        # Группируем данные по осям
        deviations_x_by_height = group_by_height(rows_x, height_tolerance)
        deviations_y_by_height = group_by_height(rows_y, height_tolerance)
        
        # Получаем все уникальные высоты
        all_heights = set(deviations_x_by_height.keys()) | set(deviations_y_by_height.keys())
        all_heights = sorted(all_heights)
        
        if not all_heights:
            return []
        
        # Формируем результат
        result = []
        for section_num, height in enumerate(all_heights):
            dev_x = deviations_x_by_height.get(height, 0.0)
            dev_y = deviations_y_by_height.get(height, 0.0)
            total_dev = np.sqrt(dev_x**2 + dev_y**2)
            
            result.append({
                'section_num': section_num,
                'height': height,
                'deviation_x': float(dev_x),
                'deviation_y': float(dev_y),
                'total_deviation': float(total_dev)
            })
        
        return result
    
    @staticmethod
    def _generate_verticality_conclusion(vertical_check: dict) -> str:
        """
        Генерирует текст заключения об отклонениях от вертикали на основе данных проверки.
        
        Args:
            vertical_check: Словарь с результатами проверки вертикальности
                {
                    'failed': int - количество превышений,
                    ...
                }
        
        Returns:
            Форматированный HTML-текст заключения
        """
        failed = vertical_check.get('failed', 0)
        
        # Определяем, превышают ли отклонения допуски
        exceeds_text = "превышают" if failed > 0 else "не превышают"
        
        conclusion_text = f"""
        <b>Заключение:</b>
        <br/><br/>
        1. Отклонения ствола от вертикали {exceeds_text} допусков СП 70.13330.2012 «Несущие и ограждающие конструкции. Актуализированная редакция СНиП 3.03.01-87» (табл.4.15) (0,001H).
        <br/><br/>
        2. Зафиксированные отклонения ствола от вертикали не препятствуют нормальной эксплуатации опоры.
        """
        
        return conclusion_text
    
    @staticmethod
    def _generate_straightness_conclusion(straightness_data: dict) -> str:
        """
        Генерирует текст заключения о стрелах прогиба на основе данных прямолинейности.
        
        Args:
            straightness_data: Словарь с данными прямолинейности
                {
                    part_num: {
                        'belts': {
                            belt_num: [
                                {
                                    'height': float,
                                    'deflection': float,  # мм
                                    'tolerance': float  # мм
                                },
                                ...
                            ]
                        }
                    }
                }
        
        Returns:
            Форматированный HTML-текст заключения
        """
        if not straightness_data:
            return ""
        
        # Находим максимальное значение прогиба по всем поясам всех частей
        max_deflection = 0.0
        max_height = 0.0
        max_belt_num = None
        max_tolerance = 0.0
        
        for part_num, part_info in straightness_data.items():
            belts_data = part_info.get('belts', {})
            for belt_num, belt_data in belts_data.items():
                for item in belt_data:
                    deflection_abs = abs(item.get('deflection', 0))
                    if deflection_abs > max_deflection:
                        max_deflection = deflection_abs
                        max_height = item.get('height', 0)
                        max_belt_num = belt_num
                        max_tolerance = item.get('tolerance', 0)
        
        if max_deflection == 0.0:
            return ""
        
        # Определяем, превышает ли максимальное значение допустимое
        exceeds_text = "превышает" if max_deflection > max_tolerance else "не превышает"
        
        # Форматируем высоту: если она больше 0, добавляем "+"
        height_str = f"+{max_height:.1f}" if max_height >= 0 else f"{max_height:.1f}"
        
        conclusion_text = f"""
        <b>Заключение:</b>
        <br/><br/>
        Стрелы прогиба рассчитаны относительно базовой линии между нижним и верхним поясами. Значения сопоставлены с нормативом δ ≤ L / 750 (ГОСТ Р 71949-2025 «Конструкции опорные антенных сооружений объектов связи. Правила приемки работ и эксплуатации»).
        <br/><br/>
        Стрела прогиба поясов башни {exceeds_text} допустимые значения. Максимальное значение составляет {max_deflection:.1f} мм на отм. {height_str} пояса №{max_belt_num} при допустимом значении {max_tolerance:.1f} мм.
        """
        
        return conclusion_text
    
    @staticmethod
    def _prepare_matplotlib_figure(figure, *, width: float, height: float, pad: float = 1.25,
                                   label_size: int = 9, title_size: int = 11):
        """Унифицированная подготовка matplotlib.figure перед экспортом."""
        if figure is None:
            return None

        original_size = figure.get_size_inches()
        figure.set_size_inches(width, height)

        for ax in figure.axes:
            ax.tick_params(labelsize=label_size)
            if ax.title and ax.title.get_text():
                ax.title.set_fontsize(title_size)
            if ax.xaxis and ax.xaxis.label:
                ax.xaxis.label.set_fontsize(label_size)
            if ax.yaxis and ax.yaxis.label:
                ax.yaxis.label.set_fontsize(label_size)
        figure.tight_layout(pad=pad)

        return original_size
    
    @staticmethod
    def _save_figure_vector_format(figure, output_path: Path, preferred_format: str = 'emf') -> Path:
        """
        Сохраняет matplotlib figure в векторном формате (EMF, WMF или SVG).
        
        Приоритет: EMF > WMF > SVG
        На Windows пытается использовать EMF/WMF, на других системах - SVG.
        
        Args:
            figure: matplotlib figure объект
            output_path: Путь для сохранения (без расширения)
            preferred_format: Предпочтительный формат ('emf', 'wmf', 'svg')
            
        Returns:
            Path к сохраненному файлу
        """
        import platform
        
        # На Windows пытаемся использовать EMF/WMF
        if platform.system() == 'Windows' and preferred_format.lower() in ('emf', 'wmf'):
            # Сначала сохраняем в SVG (универсальный векторный формат)
            svg_path = output_path.with_suffix('.svg')
            figure.savefig(str(svg_path), format='svg', bbox_inches='tight')
            
            # Пытаемся конвертировать SVG в EMF/WMF через Windows API
            try:
                import win32com.client
                
                target_format = preferred_format.lower()
                target_path = output_path.with_suffix(f'.{target_format}')
                
                # Используем Word для конвертации SVG в EMF/WMF
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                try:
                    doc = word.Documents.Add()
                    # Вставляем SVG как изображение
                    doc.InlineShapes.AddPicture(str(svg_path))
                    # Экспортируем в EMF или WMF
                    shape = doc.InlineShapes(1)
                    if target_format == 'emf':
                        # 2 = wdFormatOriginalFormat для EMF
                        shape.Export(str(target_path), ExportFormat=2)
                    else:  # wmf
                        # 1 = wdFormatOriginalFormat для WMF
                        shape.Export(str(target_path), ExportFormat=1)
                    
                    doc.Close(SaveChanges=False)
                    
                    # Удаляем временный SVG
                    if svg_path.exists():
                        svg_path.unlink()
                    
                    logger.info(f"График сохранен в {target_format.upper()} формате: {target_path}")
                    return target_path
                finally:
                    word.Quit()
                    
            except ImportError:
                logger.warning("pywin32 не установлен. Используем SVG формат.")
                return svg_path
            except Exception as e:
                logger.warning(f"Не удалось конвертировать SVG в {preferred_format.upper()}: {e}. Используем SVG.")
                return svg_path
        else:
            # На не-Windows системах или если запрошен SVG, используем SVG
            svg_path = output_path.with_suffix('.svg')
            figure.savefig(str(svg_path), format='svg', bbox_inches='tight')
            logger.info(f"График сохранен в SVG формате: {svg_path}")
            return svg_path
    
    def _register_fonts(self):
        """Регистрация шрифтов с поддержкой кириллицы"""
        try:
            # Пытаемся использовать DejaVu Sans из matplotlib
            import matplotlib
            font_path = Path(matplotlib.__file__).parent / 'mpl-data' / 'fonts' / 'ttf' / 'DejaVuSans.ttf'
            if font_path.exists():
                pdfmetrics.registerFont(TTFont('DejaVu', str(font_path)))
                pdfmetrics.registerFont(TTFont('DejaVu-Bold', str(font_path.parent / 'DejaVuSans-Bold.ttf')))
                self.font_family = 'DejaVu'
                return
            
            # Если DejaVu не найден, используем системные шрифты
            import platform
            if platform.system() == 'Windows':
                # Стандартные шрифты Windows
                arial_path = Path('C:/Windows/Fonts/arial.ttf')
                if arial_path.exists():
                    pdfmetrics.registerFont(TTFont('Arial', str(arial_path)))
                    pdfmetrics.registerFont(TTFont('Arial-Bold', str(arial_path.parent / 'arialbd.ttf')))
                    self.font_family = 'Arial'
                    return
            
            # Fallback: используем встроенные шрифты ReportLab (Vera)
            verapath = Path(__file__).parent.parent / 'venv' / 'Lib' / 'site-packages' / 'reportlab' / 'fonts' / 'Vera.ttf'
            if verapath.exists():
                pdfmetrics.registerFont(TTFont('Vera', str(verapath)))
                pdfmetrics.registerFont(TTFont('Vera-Bold', str(verapath.parent / 'VeraBd.ttf')))
                self.font_family = 'Vera'
                return
            
            # Если ничего не найдено, используем базовые шрифты
            self.font_family = 'Helvetica'
            
        except Exception as e:
            # В случае ошибки используем базовые шрифты
            self.font_family = 'Helvetica'
        
    def _create_header(self, canvas_obj, doc):
        """Создает заголовок страницы"""
        canvas_obj.saveState()
        
        # Линия сверху
        canvas_obj.setLineWidth(0.5)
        canvas_obj.line(2*cm, self.page_height - 2*cm, 
                       self.page_width - 2*cm, self.page_height - 2*cm)
        
        # Заголовок
        bold_font = f'{self.font_family}-Bold' if self.font_family in ['DejaVu', 'Arial', 'Vera'] else 'Helvetica-Bold'
        canvas_obj.setFont(bold_font, 10)
        canvas_obj.drawString(2*cm, self.page_height - 1.5*cm, 
                             'ОТЧЕТ ПО ГЕОДЕЗИЧЕСКОМУ КОНТРОЛЮ')
        
        # Дата справа
        canvas_obj.setFont(self.font_family, 9)
        date_str = self.timestamp.strftime('%d.%m.%Y')
        canvas_obj.drawRightString(self.page_width - 2*cm, 
                                   self.page_height - 1.5*cm, 
                                   f'Дата: {date_str}')
        
        canvas_obj.restoreState()
        
    def _create_footer(self, canvas_obj, doc):
        """Создает подвал страницы"""
        canvas_obj.saveState()
        
        # Линия снизу
        canvas_obj.setLineWidth(0.5)
        canvas_obj.line(2*cm, 2*cm, self.page_width - 2*cm, 2*cm)
        
        # Номер страницы
        canvas_obj.setFont(self.font_family, 9)
        page_num = canvas_obj.getPageNumber()
        text = f"Страница {page_num}"
        canvas_obj.drawCentredString(self.page_width / 2, 1.5*cm, text)
        
        # Программа слева
        canvas_obj.drawString(2*cm, 1.5*cm, 'GeoVertical Analyzer v1.0')
        
        canvas_obj.restoreState()
        
    def generate_professional_pdf(self,
                                  raw_data: pd.DataFrame,
                                  processed_data: Dict,
                                  output_path: str,
                                  project_name: str = "Объект контроля",
                                  organization: str = "",
                                  vertical_plot_widget=None,
                                  straightness_plot_widget=None,
                                  angular_measurements: Optional[Dict] = None):
        """
        Генерирует профессиональный PDF отчет
        
        Args:
            raw_data: Исходные данные
            processed_data: Результаты расчетов
            output_path: Путь для сохранения
            project_name: Название объекта
            organization: Организация
            vertical_plot_widget: Виджет графика вертикальности
            straightness_plot_widget: Виджет графика прямолинейности
        """
        try:
            # Создаем директорию для временных файлов рядом с output_path
            output_dir = Path(output_path).parent
            temp_dir = output_dir / 'temp_charts'
            temp_dir.mkdir(exist_ok=True)
            temp_files = []  # Список временных файлов для удаления

            if angular_measurements is None:
                angular_measurements = {'x': [], 'y': []}
            
            # Создаем документ
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2.5*cm,
                bottomMargin=2.5*cm
            )
            
            # Элементы документа
            story = []
            
            # Стили
            styles = getSampleStyleSheet()
            
            # Пользовательские стили
            bold_font_name = f'{self.font_family}-Bold' if self.font_family in ['DejaVu', 'Arial', 'Vera'] else 'Helvetica-Bold'
            font_name = self.font_family if self.font_family in ['DejaVu', 'Arial', 'Vera'] else 'Helvetica'
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                spaceBefore=0,
                alignment=TA_CENTER,
                fontName=bold_font_name
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=20,
                fontName=bold_font_name
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=6,
                alignment=TA_JUSTIFY,
                fontName=font_name
            )

            # === ТИТУЛЬНЫЙ ЛИСТ ===
            story.append(Spacer(1, 3*cm))
            
            # Название
            story.append(Paragraph(
                '<b>ОТЧЕТ</b><br/>по результатам геодезического контроля<br/>антенно-мачтового сооружения',
                title_style
            ))
            story.append(Spacer(1, 2*cm))
            
            # Информация об объекте
            info_data = [
                ['Объект:', project_name],
                ['Дата обследования:', self.timestamp.strftime('%d.%m.%Y')],
                ['Время:', self.timestamp.strftime('%H:%M')],
            ]
            
            if organization:
                info_data.append(['Организация:', organization])
            
            info_data.extend([
                ['Количество точек:', str(len(raw_data))],
                ['Количество поясов:', str(len(processed_data['centers']))],
                ['Программа:', 'GeoVertical Analyzer v1.0'],
            ])
            
            info_table = Table(info_data, colWidths=[5*cm, 10*cm])
            info_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 10),
                ('FONT', (0, 0), (0, -1), bold_font_name, 10),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#333333')),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            
            story.append(info_table)
            story.append(PageBreak())
            
            # === НОРМАТИВНАЯ БАЗА ===
            story.append(Paragraph('Нормативная база', heading_style))
            
            story.append(Paragraph(
                'Геодезический контроль выполнен в соответствии с требованиями:',
                body_style
            ))
            story.append(Spacer(1, 6))
            
            normatives_data = [
                ['СП 70.13330.2012', 'Несущие и ограждающие конструкции.\n'
                 'Актуализированная редакция СНиП 3.03.01-87'],
                ['', '<b>Допуск вертикальности:</b> d ≤ 0,001 × h\n'
                 'где h - высота точки от основания (м)'],
                ['', ''],
                ['ГОСТ Р 71949-2025\nКонструкции опорные\nантенных сооружений\nобъектов связи', 
                 'ГОСТ Р 71949-2025. Конструкции опорные антенных сооружений объектов связи.\n'
                 'Общие технические требования'],
                ['', '<b>Допуск прямолинейности:</b> δ ≤ L / 750\n'
                 'где L - длина секции (м)'],
            ]
            
            norm_table = Table(normatives_data, colWidths=[4*cm, 13*cm])
            norm_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), font_name, 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('LEFTPADDING', (0, 0), (0, -1), 0),
                ('RIGHTPADDING', (0, 0), (0, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            
            story.append(norm_table)
            story.append(Spacer(1, 20))

            # === ЖУРНАЛ УГЛОВЫХ ИЗМЕРЕНИЙ ===
            story.append(Paragraph('Журнал угловых измерений', heading_style))
            story.append(Paragraph(
                'Результаты теодолитных наблюдений представлены раздельно по осям X и Y. '
                'Для каждой секции приводятся чтения по левому и правому поясам, средние значения '
                'и рассчитанные отклонения.',
                body_style
            ))
            story.append(Spacer(1, 8))

            header_angular = ['№', 'Секция', 'H, м', 'По\nяс', 'KL', 'KR', 'KL–KR (″)', 'βизм', 'Bизм', 'Δβ', 'Δb, мм']
            angular_col_widths = [0.8*cm, 1.5*cm, 1.2*cm, 1.8*cm, 1.8*cm, 1.8*cm, 2.0*cm, 1.8*cm, 1.8*cm, 1.8*cm, 1.8*cm]

            def append_angular_table(axis_label: str, rows: list, ordinal: str):
                story.append(Paragraph(f'<b>Ось {axis_label}</b>', body_style))
                story.append(Spacer(1, 6))
                if not rows:
                    story.append(Paragraph('<i>Данные отсутствуют</i>', body_style))
                    story.append(Spacer(1, 6))
                    return

                table_data = [header_angular]
                for idx, row in enumerate(rows, start=1):
                    height = row.get('height')
                    height_str = f"{float(height):.3f}" if height is not None else '—'
                    belt_value = row.get('belt', '—')
                    belt_str = str(belt_value) if belt_value is not None else '—'
                    table_data.append([
                        str(idx),
                        str(row.get('section_label', '—')),
                        height_str,
                        belt_str,
                        row.get('kl_str', '—'),
                        row.get('kr_str', '—'),
                        row.get('diff_str', '—'),
                        row.get('beta_str', '—'),
                        row.get('center_str', '—'),
                        row.get('delta_str', '—'),
                        row.get('delta_mm_str', '—'),
                    ])

                angular_table = Table(table_data, colWidths=angular_col_widths, repeatRows=1)
                angular_table.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, 0), bold_font_name, 8),
                    ('FONT', (0, 1), (-1, -1), font_name, 7),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                     [colors.white, colors.HexColor('#f4f6f7')]),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                ]))

                story.append(angular_table)
                story.append(Spacer(1, 10))

            append_angular_table('X', angular_measurements.get('x', []), '1')
            append_angular_table('Y', angular_measurements.get('y', []), '2')
            story.append(Spacer(1, 14))
            
            # === ТАБЛИЦА ОТКЛОНЕНИЙ СТВОЛА ОТ ВЕРТИКАЛИ ===
            story.append(Paragraph('Таблица отклонений ствола от вертикали', heading_style))
            
            centers = processed_data['centers']
            
            story.append(Paragraph(
                f'Для {len(centers)} секций башни приведены отклонения центров от вычисленной '
                f'вертикальной оси в локальной системе координат. Таблица сопоставлена с '
                f'нормативными допусками по СП 70.13330.2012.',
                body_style
            ))
            
            results_header = [['№ секции', 'Высота,\nм', 'Откл. X,\nмм', 'Откл. Y,\nмм', 
                             'Суммарное\nоткл., мм', 'Допуск,\nмм', 'Статус']]
            results_rows = []
            
            from core.normatives import get_vertical_tolerance
            
            # Приоритет: данные из угловых измерений, затем из виджета вертикальности, затем стандартный способ
            verticality_data_from_angular = None
            if angular_measurements and (angular_measurements.get('x') or angular_measurements.get('y')):
                try:
                    verticality_data_from_angular = self._aggregate_angular_measurements_by_sections(angular_measurements)
                except Exception as e:
                    logger.warning(f"Ошибка агрегации данных угловых измерений: {e}")
                    verticality_data_from_angular = None
            
            if verticality_data_from_angular:
                # Используем данные из угловых измерений
                for item in verticality_data_from_angular:
                    height = item.get('height', 0)
                    dev_x = item.get('deviation_x', 0)
                    dev_y = item.get('deviation_y', 0)
                    total_dev = item.get('total_deviation', 0)
                    tolerance_mm = get_vertical_tolerance(height) * 1000
                    
                    status = '✓ Норма' if abs(total_dev) <= tolerance_mm else '✗ Превышение'
                    
                    results_rows.append([
                        str(item.get('section_num', 0)),
                        f"{height:.1f}",
                        f"{dev_x:+.2f}",
                        f"{dev_y:+.2f}",
                        f"{total_dev:+.2f}",
                        f"{tolerance_mm:.2f}",
                        status
                    ])
            elif vertical_plot_widget and hasattr(vertical_plot_widget, 'get_table_data'):
                try:
                    verticality_data = vertical_plot_widget.get_table_data()
                    for i, item in enumerate(verticality_data, start=1):
                        height = item.get('height', 0)
                        dev_x = item.get('deviation_x', item.get('deviation', 0))
                        dev_y = item.get('deviation_y', 0)
                        total_dev = item.get('total_deviation', item.get('deviation', 0))
                        tolerance_mm = item.get('tolerance', get_vertical_tolerance(height) * 1000)
                        
                        status = '✓ Норма' if abs(total_dev) <= tolerance_mm else '✗ Превышение'
                        
                        results_rows.append([
                            str(item.get('section_num', i)),
                            f"{height:.1f}",
                            f"{dev_x:+.2f}",
                            f"{dev_y:+.2f}",
                            f"{total_dev:+.2f}",
                            f"{tolerance_mm:.2f}",
                            status
                        ])
                except Exception as e:
                    logger.warning(f"Ошибка получения данных из виджета вертикальности: {e}")
                    # Fallback на стандартный способ
                    for i, (idx, row) in enumerate(centers.iterrows(), start=1):
                        dev_mm = row.get('deviation', 0) * 1000
                        tolerance_mm = get_vertical_tolerance(row['z']) * 1000
                        status = '✓ Норма' if abs(dev_mm) <= tolerance_mm else '✗ Превышение'
                        
                        results_rows.append([
                            str(i),
                            f"{row['z']:.1f}",
                            "—",
                            "—",
                            f"{dev_mm:.1f}",
                            f"{tolerance_mm:.1f}",
                            status
                        ])
            else:
                # Стандартный способ
                for i, (idx, row) in enumerate(centers.iterrows(), start=1):
                    dev_mm = row.get('deviation', 0) * 1000
                    tolerance_mm = get_vertical_tolerance(row['z']) * 1000
                    status = '✓ Норма' if abs(dev_mm) <= tolerance_mm else '✗ Превышение'
                    
                    results_rows.append([
                        str(i),
                        f"{row['z']:.1f}",
                        "—",
                        "—",
                        f"{dev_mm:.1f}",
                        f"{tolerance_mm:.1f}",
                        status
                    ])
            
            results_table = Table(results_header + results_rows,
                                colWidths=[1.5*cm, 2*cm, 2*cm, 2*cm, 2.5*cm, 2*cm, 2.5*cm])
            
            # Стиль с цветовым выделением
            table_style = [
                ('FONT', (0, 0), (-1, 0), bold_font_name, 9),
                ('FONT', (0, 1), (-1, -1), font_name, 8),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
            ]
            
            # Выделяем превышения красным
            for i, row in enumerate(results_rows, start=1):
                if '✗' in row[6]:
                    table_style.append(('BACKGROUND', (0, i), (-1, i), 
                                       colors.HexColor('#ffebee')))
                    table_style.append(('TEXTCOLOR', (6, i), (6, i), 
                                       colors.HexColor('#c62828')))
                else:
                    table_style.append(('TEXTCOLOR', (6, i), (6, i), 
                                       colors.HexColor('#2e7d32')))
            
            results_table.setStyle(TableStyle(table_style))
            story.append(results_table)
            story.append(Spacer(1, 10))
            
            # Вычисляем vertical_check для использования в заключении и статистике
            from core.normatives import NormativeChecker
            checker = NormativeChecker()
            
            # Приоритет: данные из угловых измерений, затем из виджета вертикальности, затем стандартный способ
            if verticality_data_from_angular:
                # Используем данные из угловых измерений
                deviations_m = []
                heights_m = []
                for item in verticality_data_from_angular:
                    height = item.get('height', 0)
                    total_dev_mm = item.get('total_deviation', 0)
                    total_dev_m = total_dev_mm / 1000.0  # Переводим из мм в метры
                    deviations_m.append(total_dev_m)
                    heights_m.append(height)
                
                if deviations_m:
                    vertical_check = checker.check_vertical_deviations(deviations_m, heights_m)
                else:
                    vertical_check = checker.check_vertical_deviations(
                        centers['deviation'].tolist(), centers['z'].tolist()
                    )
            elif vertical_plot_widget and hasattr(vertical_plot_widget, 'get_table_data'):
                try:
                    verticality_data = vertical_plot_widget.get_table_data()
                    if verticality_data:
                        deviations_m = []
                        heights_m = []
                        for item in verticality_data:
                            height = item.get('height', 0)
                            total_dev_mm = item.get('total_deviation', item.get('deviation', 0))
                            total_dev_m = total_dev_mm / 1000.0  # Переводим из мм в метры
                            deviations_m.append(total_dev_m)
                            heights_m.append(height)
                        
                        if deviations_m:
                            vertical_check = checker.check_vertical_deviations(deviations_m, heights_m)
                        else:
                            vertical_check = checker.check_vertical_deviations(
                                centers['deviation'].tolist(), centers['z'].tolist()
                            )
                    else:
                        vertical_check = checker.check_vertical_deviations(
                            centers['deviation'].tolist(), centers['z'].tolist()
                        )
                except Exception as e:
                    logger.warning(f"Ошибка получения данных вертикальности: {e}")
                    vertical_check = checker.check_vertical_deviations(
                        centers['deviation'].tolist(), centers['z'].tolist()
                    )
            else:
                vertical_check = checker.check_vertical_deviations(
                    centers['deviation'].tolist(), centers['z'].tolist()
                )
            
            # График вертикальности сразу после таблицы
            if vertical_plot_widget and hasattr(vertical_plot_widget, 'figure'):
                story.append(Paragraph('<b>График вертикальности</b>', body_style))
                story.append(Spacer(1, 6))
                
                fig = vertical_plot_widget.figure
                original_size = self._prepare_matplotlib_figure(fig, width=7.8, height=6.2, pad=1.5)
                
                tmp_path_svg = temp_dir / 'vertical_plot.svg'
                fig.savefig(str(tmp_path_svg), format='svg', bbox_inches='tight')
                temp_files.append(str(tmp_path_svg))
                
                tmp_path_pdf = temp_dir / 'vertical_plot.pdf'
                fig.savefig(str(tmp_path_pdf), format='pdf', bbox_inches='tight')
                temp_files.append(str(tmp_path_pdf))
                
                tmp_path = temp_dir / 'vertical_plot.png'
                fig.savefig(str(tmp_path), dpi=320, bbox_inches='tight')
                temp_files.append(str(tmp_path))
                img = Image(str(tmp_path), width=16*cm, height=12*cm)
                
                story.append(img)
                story.append(Spacer(1, 6))
                story.append(Paragraph(
                    '<i>Рис. 1. Отклонения центров секций от вертикальной оси мачты</i>',
                    ParagraphStyle('Caption', parent=body_style, 
                                  fontSize=9, alignment=TA_CENTER)
                ))
                story.append(Spacer(1, 20))
                
                if original_size is not None:
                    fig.set_size_inches(original_size)
                
                # Заключение после графика вертикальности
                conclusion_text = self._generate_verticality_conclusion(vertical_check)
                story.append(Paragraph(conclusion_text, body_style))
                story.append(Spacer(1, 10))
            else:
                # Если графика нет, но нужно заключение, вычисляем vertical_check
                if 'vertical_check' not in locals():
                    from core.normatives import NormativeChecker
                    checker = NormativeChecker()
                    vertical_check = checker.check_vertical_deviations(
                        centers['deviation'].tolist(), centers['z'].tolist()
                    )
                
                # Заключение даже если нет графика
                conclusion_text = self._generate_verticality_conclusion(vertical_check)
                story.append(Paragraph(conclusion_text, body_style))
                story.append(Spacer(1, 10))
            
            # Таблица прогибов (прямолинейность)
            story.append(Paragraph('Таблица стрел прогиба поясов ствола (прямолинейность)', heading_style))
            story.append(Paragraph(
                'Стрелы прогиба рассчитаны относительно базовой линии между нижним и верхним поясами. '
                'Значения сравниваются с нормативом δ ≤ L / 750.',
                body_style
            ))
            story.append(Spacer(1, 8))

            if straightness_plot_widget and hasattr(straightness_plot_widget, 'get_all_belts_data'):
                try:
                    straightness_data = straightness_plot_widget.get_all_belts_data()
                    
                    if straightness_data:
                        # Новая структура: данные сгруппированы по частям
                        # straightness_data = {part_num: {'min_height': float, 'max_height': float, 'belts': {belt_num: [data]}}}
                        
                        sorted_parts = sorted(straightness_data.keys())
                        
                        # Общий счетчик рисунков для всех частей (начинается с 2, т.к. рис. 1 - вертикальность)
                        global_figure_index = 2
                        
                        for part_num in sorted_parts:
                            part_info = straightness_data[part_num]
                            min_height = part_info.get('min_height', 0.0)
                            max_height = part_info.get('max_height', 0.0)
                            belts_data = part_info.get('belts', {})
                            
                            if not belts_data:
                                continue
                            
                            # Заголовок для части (формат высот: "0,0 - 20,800")
                            # Форматируем высоты: одна цифра после запятой для min, три для max
                            min_str = f"{min_height:.1f}".replace('.', ',')
                            max_str = f"{max_height:.3f}".replace('.', ',')
                            part_title = f"Часть {part_num} ({min_str} - {max_str})"
                            story.append(Paragraph(f'<b>4.{part_num if part_num > 1 else ""}. {part_title}</b>', heading_style if part_num == 1 else body_style))
                            story.append(Spacer(1, 6))
                            
                            # Получаем все уникальные высоты для этой части
                            all_heights = set()
                            max_tolerance = 0
                            for belt_data in belts_data.values():
                                for item in belt_data:
                                    all_heights.add(round(item['height'], 1))
                                    max_tolerance = max(max_tolerance, item.get('tolerance', 0))
                            
                            sorted_heights = sorted(all_heights)
                            sorted_belts = sorted(belts_data.keys())
                            
                            # Создаем словарь для быстрого доступа
                            belt_height_deflection = {}
                            for belt_num, belt_data in belts_data.items():
                                for item in belt_data:
                                    height_rounded = round(item['height'], 1)
                                    belt_height_deflection[(belt_num, height_rounded)] = item.get('deflection', 0)
                            
                            # Заголовок таблицы
                            deflection_header = [['Высота,\nм']]
                            for belt_num in sorted_belts:
                                deflection_header[0].append(f'Пояс {belt_num},\nмм')
                            deflection_header[0].append('Допуск,\nмм')
                            
                            # Строки таблицы
                            deflection_rows = []
                            for height in sorted_heights:
                                row = [f"{height:.1f}"]
                                for belt_num in sorted_belts:
                                    key = (belt_num, height)
                                    if key in belt_height_deflection:
                                        deflection = belt_height_deflection[key]
                                        row.append(f"{deflection:+.1f}")
                                    else:
                                        row.append("—")
                                row.append(f"±{max_tolerance:.1f}")
                                deflection_rows.append(row)
                            
                            # Создаем таблицу
                            col_widths = [2*cm] + [2*cm] * len(sorted_belts) + [2*cm]
                            deflection_table = Table(deflection_header + deflection_rows, colWidths=col_widths)
                            
                            # Стиль таблицы прогибов
                            deflection_style = [
                                ('FONT', (0, 0), (-1, 0), bold_font_name, 9),
                                ('FONT', (0, 1), (-1, -1), font_name, 8),
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                ('TOPPADDING', (0, 0), (-1, -1), 4),
                            ]
                            
                            # Выделяем превышения
                            for i, row in enumerate(deflection_rows, start=1):
                                for j in range(1, len(row) - 1):  # Пропускаем первый (высота) и последний (допуск)
                                    try:
                                        deflection_val = float(row[j].replace('+', '').replace('—', '0'))
                                        if abs(deflection_val) > max_tolerance:
                                            deflection_style.append(('BACKGROUND', (j, i), (j, i), 
                                                                   colors.HexColor('#ffebee')))
                                            deflection_style.append(('TEXTCOLOR', (j, i), (j, i), 
                                                                   colors.HexColor('#c62828')))
                                    except ValueError:
                                        pass
                            
                            deflection_table.setStyle(TableStyle(deflection_style))
                            story.append(deflection_table)
                            story.append(Spacer(1, 10))
                            
                            # Графики прямолинейности для этой части сразу после таблицы
                            story.append(Paragraph('<b>Графики прямолинейности</b>', body_style))
                            story.append(Spacer(1, 6))
                            
                            # Получаем графики для этой части
                            part_figures = []
                            if hasattr(straightness_plot_widget, 'get_part_figures_for_pdf'):
                                try:
                                    part_figures = straightness_plot_widget.get_part_figures_for_pdf(part_num, group_size=2)
                                except Exception as exc:  # noqa: BLE001
                                    logger.warning(f"Не удалось получить графики для части {part_num}: {exc}")
                            
                            if not part_figures:
                                # Fallback: пробуем получить все графики и отфильтровать
                                try:
                                    if hasattr(straightness_plot_widget, 'get_grouped_figures_for_pdf'):
                                        all_figures = straightness_plot_widget.get_grouped_figures_for_pdf()
                                        # Фильтруем по поясам, принадлежащим этой части
                                        part_belts = set(sorted_belts)
                                        for belt_group, figure in all_figures:
                                            if any(belt in part_belts for belt in belt_group):
                                                part_figures.append((belt_group, figure))
                                except Exception as exc:  # noqa: BLE001
                                    logger.warning(f"Не удалось получить графики (fallback) для части {part_num}: {exc}")
                            
                            if part_figures:
                                for belt_group, figure in part_figures:
                                    width = 11.0 if len(belt_group) > 1 else 8.5
                                    original_size = self._prepare_matplotlib_figure(
                                        figure,
                                        width=width,
                                        height=5.8,
                                        pad=1.6,
                                        label_size=9,
                                        title_size=11
                                    )
                                    tmp_path = temp_dir / f'straightness_part_{part_num}_group_{global_figure_index}.png'
                                    figure.savefig(str(tmp_path), dpi=320, bbox_inches='tight')
                                    temp_files.append(str(tmp_path))
                                    img = Image(str(tmp_path), width=16*cm, height=11*cm)
                                    story.append(img)
                                    story.append(Spacer(1, 6))

                                    belts_caption = ', '.join(str(b) for b in belt_group)
                                    story.append(Paragraph(
                                        f'<i>Рис. {global_figure_index}. Отклонения от прямолинейности по поясам {belts_caption}</i>',
                                        ParagraphStyle('Caption', parent=body_style,
                                                      fontSize=9, alignment=TA_CENTER)
                                    ))
                                    story.append(Spacer(1, 14))
                                    global_figure_index += 1

                                    if original_size is not None:
                                        figure.set_size_inches(original_size)
                            else:
                                story.append(Paragraph('<i>Графики для этой части недоступны</i>', body_style))
                            
                            story.append(Spacer(1, 10))  # Отступ после графиков части
                        
                except Exception as e:
                    logger.error(f"Ошибка генерации таблицы прогибов: {e}")
            else:
                story.append(Paragraph('<i>Данные прямолинейности отсутствуют</i>', body_style))
                story.append(Spacer(1, 10))
            
            # Заключение о прямолинейности после всех графиков
            if straightness_plot_widget and hasattr(straightness_plot_widget, 'get_all_belts_data'):
                try:
                    straightness_data_for_conclusion = straightness_plot_widget.get_all_belts_data()
                    if straightness_data_for_conclusion:
                        conclusion_text_straightness = self._generate_straightness_conclusion(straightness_data_for_conclusion)
                        if conclusion_text_straightness:
                            story.append(Paragraph(conclusion_text_straightness, body_style))
                            story.append(Spacer(1, 10))
                except Exception as e:
                    logger.warning(f"Ошибка формирования заключения о прямолинейности: {e}")

            # Статистика проверки вертикальности (используем уже вычисленный vertical_check)
            # vertical_check был вычислен ранее после таблицы отклонений
            stats_text = f"""
            <b>Статистика проверки вертикальности:</b><br/>
            • Всего секций: {vertical_check['total']}<br/>
            • В пределах нормы: {vertical_check['passed']} 
              ({vertical_check['passed']/vertical_check['total']*100:.1f}%)<br/>
            • Превышение допуска: {vertical_check['failed']} 
              ({vertical_check['failed']/vertical_check['total']*100:.1f}%)
            """
            
            story.append(Paragraph(stats_text, body_style))
            
            if vertical_check['non_compliant']:
                story.append(Spacer(1, 10))
                story.append(Paragraph(
                    '<b>Секции с превышением допуска:</b>',
                    body_style
                ))
                
                for item in vertical_check['non_compliant']:
                    # excess уже в метрах, переводим в миллиметры
                    excess_mm = item['excess'] * 1000
                    story.append(Paragraph(
                        f"• Секция {item['index']+1} (высота {item['height']:.1f}м): "
                        f"превышение {excess_mm:.2f} мм",
                        body_style
                    ))
            
            # Генерируем PDF
            doc.build(story, onFirstPage=self._create_header, 
                     onLaterPages=self._create_header)
            
            # Удаляем временные файлы после генерации
            for temp_file in temp_files:
                try:
                    if Path(temp_file).exists():
                        Path(temp_file).unlink()
                except Exception as e:
                    print(f"Не удалось удалить временный файл {temp_file}: {e}")
            
            # Удаляем директорию, если она пуста
            try:
                if temp_dir.exists() and not any(temp_dir.iterdir()):
                    temp_dir.rmdir()
            except Exception as e:
                print(f"Не удалось удалить временную директорию {temp_dir}: {e}")
            
        except Exception as e:
            raise ValueError(f"Ошибка генерации профессионального PDF отчета: {str(e)}")
    
    def generate_professional_docx(self,
                                  raw_data: pd.DataFrame,
                                  processed_data: Dict,
                                  output_path: str,
                                  project_name: str = "Объект контроля",
                                  organization: str = "",
                                  verticality_widget=None,
                                  straightness_widget=None,
                                  angular_measurements: Optional[Dict] = None):
        """
        Генерирует профессиональный DOCX отчет, идентичный PDF
        
        Args:
            raw_data: Исходные данные
            processed_data: Результаты расчетов
            output_path: Путь для сохранения
            project_name: Название объекта
            organization: Организация
            verticality_widget: Виджет графика вертикальности
            straightness_widget: Виджет графика прямолинейности
        """
        try:
            from docx import Document
            from docx.shared import Cm, Inches, Pt, RGBColor
            from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from pathlib import Path
            import numpy as np
            from core.normatives import get_vertical_tolerance
            
            if angular_measurements is None:
                angular_measurements = {'x': [], 'y': []}

            centers = processed_data['centers']
            
            # Создаем директорию для временных файлов
            output_dir = Path(output_path).parent
            temp_dir = output_dir / 'temp_charts'
            temp_dir.mkdir(exist_ok=True)
            temp_files = []
            
            # Создаем документ
            doc = Document()
            
            # Настройка полей страницы: левое 2,0 см, правое 0,75 см
            section = doc.sections[0]
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(0.75)
            
            def format_table(table, *, header_rows: int = 1, font_size: int = 9, column_widths=None):
                """Единообразное форматирование таблиц в DOCX-отчете."""
                if table is None:
                    return

                # Устанавливаем стиль таблицы с границами
                table.style = 'Table Grid'

                table.autofit = True
                table.allow_autofit = True

                if column_widths:
                    for col_idx, width_cm in enumerate(column_widths):
                        if col_idx >= len(table.columns):
                            break
                        for cell in table.columns[col_idx].cells:
                            cell.width = Cm(width_cm)

                for row_idx, row in enumerate(table.rows):
                    row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.15
                            if paragraph.alignment is None:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                if run.font.size is None:
                                    run.font.size = Pt(font_size)
                        if row_idx < header_rows and cell.paragraphs:
                            for run in cell.paragraphs[0].runs:
                                run.font.bold = True
            
            # Настройка шрифта по умолчанию
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)
            
            # === ТИТУЛЬНЫЙ ЛИСТ ===
            doc.add_paragraph()  # Пустой абзац сверху
            doc.add_paragraph()  # Еще один
            doc.add_paragraph()  # Еще один
            
            # Заголовок
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run('ОТЧЕТ')
            run.font.bold = True
            run.font.size = Pt(16)
            
            title2 = doc.add_paragraph()
            title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = title2.add_run('по результатам геодезического контроля')
            run2.font.size = Pt(16)
            
            title3 = doc.add_paragraph()
            title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = title3.add_run('антенно-мачтового сооружения')
            run3.font.size = Pt(16)
            
            doc.add_paragraph()  # Отступ
            doc.add_paragraph()  # Отступ
            
            # Информация об объекте
            info_data = [
                ['Объект:', project_name],
                ['Дата обследования:', self.timestamp.strftime('%d.%m.%Y')],
                ['Время:', self.timestamp.strftime('%H:%M')],
            ]
            
            if organization:
                info_data.append(['Организация:', organization])
            
            info_data.extend([
                ['Количество точек:', str(len(raw_data))],
                ['Количество поясов:', str(len(processed_data['centers']))],
                ['Программа:', 'GeoVertical Analyzer v1.0'],
            ])
            
            info_table = doc.add_table(rows=len(info_data), cols=2)
            info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, row_data in enumerate(info_data):
                row = info_table.rows[i]
                row.cells[0].text = row_data[0]
                row.cells[1].text = row_data[1]
                
                # Форматирование
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            
            format_table(
                info_table,
                header_rows=0,
                font_size=10,
                column_widths=[4.0, 9.0]
            )
            
            doc.add_page_break()
            
            # === НОРМАТИВНАЯ БАЗА ===
            heading = doc.add_paragraph('Нормативная база')
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            text = doc.add_paragraph('Геодезический контроль выполнен в соответствии с требованиями:')
            text.runs[0].font.size = Pt(10)
            
            # Таблица нормативов
            norm_table = doc.add_table(rows=5, cols=2)
            norm_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            norm_table.rows[0].cells[0].text = 'СП 70.13330.2012'
            norm_table.rows[0].cells[1].text = 'Несущие и ограждающие конструкции.\nАктуализированная редакция СНиП 3.03.01-87'
            
            norm_table.rows[1].cells[0].text = ''
            norm_bold = norm_table.rows[1].cells[1].paragraphs[0]
            norm_run1 = norm_bold.add_run('Допуск вертикальности: ')
            norm_run1.font.bold = True
            norm_run2 = norm_bold.add_run('d ≤ 0,001 × h\nгде h - высота точки от основания (м)')
            
            norm_table.rows[2].cells[0].text = ''
            norm_table.rows[2].cells[1].text = ''
            
            norm_table.rows[3].cells[0].text = 'ГОСТ Р 71949-2025\nКонструкции опорные\nантенных сооружений\nобъектов связи'
            norm_table.rows[3].cells[1].text = 'ГОСТ Р 71949-2025. Конструкции опорные антенных сооружений объектов связи.\nОбщие технические требования'
            
            norm_table.rows[4].cells[0].text = ''
            norm_bold2 = norm_table.rows[4].cells[1].paragraphs[0]
            norm_run3 = norm_bold2.add_run('Допуск прямолинейности: ')
            norm_run3.font.bold = True
            norm_run4 = norm_bold2.add_run('δ ≤ L / 750\nгде L - длина секции (м)')
            
            # Форматирование таблицы нормативов
            for row in norm_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            format_table(
                norm_table,
                header_rows=0,
                font_size=9,
                column_widths=[4.2, 10.8]
            )
            
            doc.add_paragraph()  # Отступ
            
            doc.add_page_break()
            
            # === ЖУРНАЛ УГЛОВЫХ ИЗМЕРЕНИЙ ===
            heading_angles = doc.add_paragraph('Журнал угловых измерений')
            heading_angles.runs[0].font.bold = True
            heading_angles.runs[0].font.size = Pt(12)

            angles_desc = doc.add_paragraph(
                'Результаты теодолитных наблюдений по видимым поясам башни. '
                'Показаны чтения по левому и правому поясам, средние значения и рассчитанные отклонения.'
            )
            angles_desc.runs[0].font.size = Pt(10)

            headers_angular = ['№', 'Секция', 'H, м', 'Пояс', 'KL', 'KR', 'KL–KR (″)', 'βизм', 'Bизм', 'Δβ', 'Δb, мм']
            header_angular_with_break = ['№', 'Секция', 'H, м', 'По\nяс', 'KL', 'KR', 'KL–KR (″)', 'βизм', 'Bизм', 'Δβ', 'Δb, мм']
            angular_widths = [0.8, 1.5, 1.2, 1.8, 1.8, 1.8, 2.0, 1.8, 1.8, 1.8, 1.8]

            def append_angular_table(axis_label: str, rows: list, ordinal: str):
                sub_heading = doc.add_paragraph(f'Ось {axis_label}')
                sub_heading.runs[0].font.bold = True
                sub_heading.runs[0].font.size = Pt(10)

                if not rows:
                    note = doc.add_paragraph('Данные отсутствуют')
                    note.runs[0].italic = True
                    note.runs[0].font.size = Pt(10)
                    doc.add_paragraph()
                    return

                table = doc.add_table(rows=len(rows) + 1, cols=len(headers_angular))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, width_cm in enumerate(angular_widths):
                    table.columns[i].width = Cm(width_cm)

                header_row = table.rows[0]
                for idx_header, header_text in enumerate(headers_angular):
                    # Для столбца "Пояс" используем разбитый заголовок
                    if idx_header == 3:  # Столбец "Пояс" (индекс 3)
                        header_row.cells[idx_header].text = header_angular_with_break[idx_header]
                    else:
                        header_row.cells[idx_header].text = header_text
                    paragraph = header_row.cells[idx_header].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Включаем перенос текста для столбца "Пояс"
                    if idx_header == 3:
                        paragraph.paragraph_format.wrap_text = True
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(9)
                
                # Устанавливаем высоту строки заголовка для переноса
                header_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO

                for idx, row in enumerate(rows, start=1):
                    data_row = table.rows[idx]
                    height = row.get('height')
                    height_str = f"{float(height):.1f}" if height is not None else '—'
                    belt_value = row.get('belt', '—')
                    # Преобразуем дробные значения поясов (1.0, 2.0, 3.0) в целые числа (1, 2, 3)
                    if belt_value is not None and belt_value != '—':
                        try:
                            belt_float = float(belt_value)
                            if belt_float.is_integer():
                                belt_str = str(int(belt_float))
                            else:
                                belt_str = str(belt_value)
                        except (ValueError, TypeError):
                            belt_str = str(belt_value)
                    else:
                        belt_str = '—'
                    values = [
                        str(idx),
                        str(row.get('section_label', '—')),
                        height_str,
                        belt_str,
                        row.get('kl_str', '—'),
                        row.get('kr_str', '—'),
                        row.get('diff_str', '—'),
                        row.get('beta_str', '—'),
                        row.get('center_str', '—'),
                        row.get('delta_str', '—'),
                        row.get('delta_mm_str', '—'),
                    ]

                    for col_idx, value in enumerate(values):
                        data_row.cells[col_idx].text = value
                        paragraph = data_row.cells[col_idx].paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Включаем перенос текста для столбца "Пояс"
                        if col_idx == 3:
                            paragraph.paragraph_format.wrap_text = True
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                
                format_table(
                    table,
                    header_rows=1,
                    font_size=8,
                    column_widths=angular_widths
                )

                doc.add_paragraph()

            append_angular_table('X', angular_measurements.get('x', []), '1')
            append_angular_table('Y', angular_measurements.get('y', []), '2')

            doc.add_paragraph()  # Отступ

            # === ТАБЛИЦА ОТКЛОНЕНИЙ СТВОЛА ОТ ВЕРТИКАЛИ ===
            heading3 = doc.add_paragraph('Таблица отклонений ствола от вертикали')
            heading3.runs[0].font.bold = True
            heading3.runs[0].font.size = Pt(12)

            doc.add_paragraph()  # Отступ

            # Таблица результатов вертикальности (полная таблица)
            # Приоритет: данные из угловых измерений, затем из виджета вертикальности, затем стандартный способ
            used_widget_data = False
            verticality_data_from_angular = None
            if angular_measurements and (angular_measurements.get('x') or angular_measurements.get('y')):
                try:
                    verticality_data_from_angular = self._aggregate_angular_measurements_by_sections(angular_measurements)
                except Exception as e:
                    logger.warning(f"Ошибка агрегации данных угловых измерений: {e}")
                    verticality_data_from_angular = None
            
            # Определяем количество строк и источник данных
            if verticality_data_from_angular:
                num_rows = len(verticality_data_from_angular) + 1
            elif verticality_widget and hasattr(verticality_widget, 'get_table_data'):
                try:
                    verticality_data = verticality_widget.get_table_data()
                    num_rows = len(verticality_data) + 1
                except:
                    num_rows = len(centers) + 1
            else:
                num_rows = len(centers) + 1

            results_table = doc.add_table(rows=num_rows, cols=7)
            results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Заголовок
            header_row = results_table.rows[0]
            header_row.cells[0].text = '№ секции'
            header_row.cells[1].text = 'Высота,\nм'
            header_row.cells[2].text = 'Откл. X,\nмм'
            header_row.cells[3].text = 'Откл. Y,\nмм'
            header_row.cells[4].text = 'Суммарное\nоткл., мм'
            header_row.cells[5].text = 'Допуск,\nмм'
            header_row.cells[6].text = 'Статус'

            for paragraph in header_row.cells[0].paragraphs + header_row.cells[1].paragraphs + \
                            header_row.cells[2].paragraphs + header_row.cells[3].paragraphs + \
                            header_row.cells[4].paragraphs + header_row.cells[5].paragraphs + \
                            header_row.cells[6].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].font.bold = True
                paragraph.runs[0].font.size = Pt(9)

            # Данные
            if verticality_data_from_angular:
                # Используем данные из угловых измерений
                
                for i, item in enumerate(verticality_data_from_angular):
                    if i + 1 >= num_rows:
                        break
                    data_row = results_table.rows[i + 1]
                    height = item.get('height', 0)
                    dev_x = item.get('deviation_x', 0)
                    dev_y = item.get('deviation_y', 0)
                    total_dev = item.get('total_deviation', 0)
                    tolerance_mm = get_vertical_tolerance(height) * 1000
                    
                    data_row.cells[0].text = str(item.get('section_num', i))
                    data_row.cells[1].text = f"{height:.1f}"
                    data_row.cells[2].text = f"{dev_x:+.1f}"
                    data_row.cells[3].text = f"{dev_y:+.1f}"
                    # Суммарное отклонение: округляем до десятого знака и выводим по модулю
                    data_row.cells[4].text = f"{abs(total_dev):.1f}"
                    data_row.cells[5].text = f"{tolerance_mm:.1f}"
                    
                    # Статус
                    if abs(total_dev) > tolerance_mm:
                        status_para = data_row.cells[6].paragraphs[0]
                        status_run = status_para.add_run('✗ Превышение')
                        status_run.font.color.rgb = RGBColor(255, 0, 0)
                        status_run.font.bold = True
                        data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        status_para = data_row.cells[6].paragraphs[0]
                        status_run = status_para.add_run('✓ Норма')
                        status_run.font.color.rgb = RGBColor(0, 128, 0)
                        data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Форматирование
                    for j in range(7):
                        for paragraph in data_row.cells[j].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                
                used_widget_data = True
            elif verticality_widget and hasattr(verticality_widget, 'get_table_data'):
                try:
                    verticality_data = verticality_widget.get_table_data()
                    # Заполняем таблицу данными из виджета
                    for i, item in enumerate(verticality_data):
                        if i + 1 >= num_rows:
                            break
                        data_row = results_table.rows[i + 1]
                        data_row.cells[0].text = str(item.get('section_num', i + 1))
                        data_row.cells[1].text = f"{item['height']:.1f}"
                        
                        dev_x = item.get('deviation_x', item.get('deviation', 0))
                        dev_y = item.get('deviation_y', 0)
                        total_dev = item.get('total_deviation', item.get('deviation', 0))
                        tolerance_mm = item.get('tolerance', get_vertical_tolerance(item['height']) * 1000)
                        
                        data_row.cells[2].text = f"{dev_x:+.1f}"
                        data_row.cells[3].text = f"{dev_y:+.1f}"
                        # Суммарное отклонение: округляем до десятого знака и выводим по модулю
                        data_row.cells[4].text = f"{abs(total_dev):.1f}"
                        data_row.cells[5].text = f"{tolerance_mm:.1f}"
                        
                        # Статус
                        if abs(total_dev) > tolerance_mm:
                            status_para = data_row.cells[6].paragraphs[0]
                            status_run = status_para.add_run('✗ Превышение')
                            status_run.font.color.rgb = RGBColor(255, 0, 0)
                            status_run.font.bold = True
                            data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            status_para = data_row.cells[6].paragraphs[0]
                            status_run = status_para.add_run('✓ Норма')
                            status_run.font.color.rgb = RGBColor(0, 128, 0)
                            data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Форматирование
                        for j in range(7):
                            for paragraph in data_row.cells[j].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                    
                    used_widget_data = True
                except Exception as e:
                    logger.warning(f"Ошибка получения данных из виджета вертикальности: {e}")
            
            # Вычисляем vertical_check для использования в заключении (перед графиком)
            from core.normatives import NormativeChecker
            checker = NormativeChecker()
            
            # Приоритет: данные из угловых измерений, затем из виджета вертикальности, затем стандартный способ
            if verticality_data_from_angular:
                deviations_m = []
                heights_m = []
                for item in verticality_data_from_angular:
                    height = item.get('height', 0)
                    total_dev_mm = item.get('total_deviation', 0)
                    total_dev_m = total_dev_mm / 1000.0  # Переводим из мм в метры
                    deviations_m.append(total_dev_m)
                    heights_m.append(height)
                
                if deviations_m:
                    vertical_check = checker.check_vertical_deviations(deviations_m, heights_m)
                else:
                    vertical_check = checker.check_vertical_deviations(
                        centers['deviation'].tolist(), centers['z'].tolist()
                    )
            elif verticality_widget and hasattr(verticality_widget, 'get_table_data'):
                try:
                    verticality_data = verticality_widget.get_table_data()
                    if verticality_data:
                        deviations_m = []
                        heights_m = []
                        for item in verticality_data:
                            height = item.get('height', 0)
                            total_dev_mm = item.get('total_deviation', item.get('deviation', 0))
                            total_dev_m = total_dev_mm / 1000.0  # Переводим из мм в метры
                            deviations_m.append(total_dev_m)
                            heights_m.append(height)
                        
                        if deviations_m:
                            vertical_check = checker.check_vertical_deviations(deviations_m, heights_m)
                        else:
                            vertical_check = checker.check_vertical_deviations(
                                centers['deviation'].tolist(), centers['z'].tolist()
                            )
                    else:
                        vertical_check = checker.check_vertical_deviations(
                            centers['deviation'].tolist(), centers['z'].tolist()
                        )
                except Exception as e:
                    logger.warning(f"Ошибка получения данных вертикальности: {e}")
                    vertical_check = checker.check_vertical_deviations(
                        centers['deviation'].tolist(), centers['z'].tolist()
                    )
            else:
                vertical_check = checker.check_vertical_deviations(
                    centers['deviation'].tolist(), centers['z'].tolist()
                )
            
            # График вертикальности сразу после таблицы
            if verticality_widget and hasattr(verticality_widget, 'figure'):
                doc.add_paragraph()  # Отступ
                heading3_1 = doc.add_paragraph('График вертикальности')
                heading3_1.runs[0].font.bold = True
                heading3_1.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()  # Отступ
                
                # Сохраняем график в векторном формате (EMF/WMF/SVG)
                tmp_path_base = temp_dir / 'vertical_docx'
                fig = verticality_widget.figure
                original_size = self._prepare_matplotlib_figure(fig, width=7.5, height=6.0, pad=1.4)
                
                # Сохраняем в векторном формате (предпочтительно EMF)
                tmp_path = self._save_figure_vector_format(fig, tmp_path_base, preferred_format='emf')
                temp_files.append(str(tmp_path))
                
                # Вставляем график (python-docx поддерживает EMF, WMF и SVG)
                try:
                    doc.add_picture(str(tmp_path), width=Inches(6))
                except Exception as e:
                    # Если не удалось вставить векторный формат, используем высококачественный PNG как fallback
                    logger.warning(f"Не удалось вставить векторный график: {e}. Используем PNG.")
                    png_path = temp_dir / 'vertical_docx.png'
                    fig.savefig(str(png_path), dpi=600, bbox_inches='tight')
                    temp_files.append(str(png_path))
                    doc.add_picture(str(png_path), width=Inches(6))
                
                if original_size is not None:
                    fig.set_size_inches(original_size)
                
                caption = doc.add_paragraph('Рис. 1. Отклонения центров секций от вертикальной оси мачты')
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(9)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Отступ
                
                # Заключение после графика вертикальности
                failed = vertical_check.get('failed', 0)
                exceeds_text = "превышают" if failed > 0 else "не превышают"
                
                conclusion_heading = doc.add_paragraph()
                conclusion_heading_run = conclusion_heading.add_run('Заключение:')
                conclusion_heading_run.font.bold = True
                conclusion_heading_run.font.size = Pt(10)
                
                conclusion_para1 = doc.add_paragraph()
                conclusion_text1 = f"1. Отклонения ствола от вертикали {exceeds_text} допусков СП 70.13330.2012 «Несущие и ограждающие конструкции. Актуализированная редакция СНиП 3.03.01-87» (табл.4.15) (0,001H)."
                conclusion_para1.add_run(conclusion_text1).font.size = Pt(10)
                
                conclusion_para2 = doc.add_paragraph()
                conclusion_text2 = "2. Зафиксированные отклонения ствола от вертикали не препятствуют нормальной эксплуатации опоры."
                conclusion_para2.add_run(conclusion_text2).font.size = Pt(10)
                
                doc.add_paragraph()  # Отступ
            
            else:
                # Заключение даже если нет графика (vertical_check уже вычислен ранее)
                failed = vertical_check.get('failed', 0)
                exceeds_text = "превышают" if failed > 0 else "не превышают"
                
                conclusion_heading = doc.add_paragraph()
                conclusion_heading_run = conclusion_heading.add_run('Заключение:')
                conclusion_heading_run.font.bold = True
                conclusion_heading_run.font.size = Pt(10)
                
                conclusion_para1 = doc.add_paragraph()
                conclusion_text1 = f"1. Отклонения ствола от вертикали {exceeds_text} допусков СП 70.13330.2012 «Несущие и ограждающие конструкции. Актуализированная редакция СНиП 3.03.01-87» (табл.4.15) (0,001H)."
                conclusion_para1.add_run(conclusion_text1).font.size = Pt(10)
                
                conclusion_para2 = doc.add_paragraph()
                conclusion_text2 = "2. Зафиксированные отклонения ствола от вертикали не препятствуют нормальной эксплуатации опоры."
                conclusion_para2.add_run(conclusion_text2).font.size = Pt(10)
                
                doc.add_paragraph()  # Отступ
            
            # Стандартный цикл (fallback)
            if not used_widget_data:
                for i, (idx, row) in enumerate(centers.iterrows()):
                    if i + 1 >= num_rows:
                        break
                    data_row = results_table.rows[i + 1]
                    data_row.cells[0].text = str(i + 1)
                    data_row.cells[1].text = f"{row['z']:.1f}"
                    data_row.cells[2].text = "—"
                    data_row.cells[3].text = "—"
                    
                    dev_mm = row.get('deviation', 0) * 1000
                    tolerance_mm = get_vertical_tolerance(row['z']) * 1000
                    # Суммарное отклонение: округляем до десятого знака и выводим по модулю
                    data_row.cells[4].text = f"{abs(dev_mm):.1f}"
                    data_row.cells[5].text = f"{tolerance_mm:.1f}"
                    
                    # Статус
                    if abs(dev_mm) > tolerance_mm:
                        status_para = data_row.cells[6].paragraphs[0]
                        status_run = status_para.add_run('✗ Превышение')
                        status_run.font.color.rgb = RGBColor(255, 0, 0)
                        status_run.font.bold = True
                        data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        status_para = data_row.cells[6].paragraphs[0]
                        status_run = status_para.add_run('✓ Норма')
                        status_run.font.color.rgb = RGBColor(0, 128, 0)
                        data_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Форматирование
                    for j in range(7):
                        for paragraph in data_row.cells[j].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
            
            format_table(
                results_table,
                header_rows=1,
                font_size=8,
                column_widths=[2.4, 2.4, 2.2, 2.2, 2.4, 2.2, 2.6]
            )
            
            # Таблица прогибов (прямолинейность)
            heading4 = doc.add_paragraph('Таблица стрел прогиба поясов ствола (прямолинейность)')
            heading4.runs[0].font.bold = True
            heading4.runs[0].font.size = Pt(12)

            desc4 = doc.add_paragraph('Стрелы прогиба рассчитаны относительно базовой линии между нижним и верхним поясами. Значения сопоставлены с нормативом δ ≤ L / 750.')
            desc4.runs[0].font.size = Pt(10)

            if straightness_widget and hasattr(straightness_widget, 'get_all_belts_data'):
                try:
                    straightness_data = straightness_widget.get_all_belts_data()
                    
                    if straightness_data:
                        # Новая структура: данные сгруппированы по частям
                        sorted_parts = sorted(straightness_data.keys())
                        
                        # Общий счетчик рисунков для всех частей (начинается с 2, т.к. рис. 1 - вертикальность)
                        global_figure_index = 2
                        
                        for part_num in sorted_parts:
                            part_info = straightness_data[part_num]
                            min_height = part_info.get('min_height', 0.0)
                            max_height = part_info.get('max_height', 0.0)
                            belts_data = part_info.get('belts', {})
                            
                            if not belts_data:
                                continue
                            
                            # Заголовок для части (формат высот: "0,0 - 20,800")
                            # Форматируем высоты: одна цифра после запятой для min, три для max
                            min_str = f"{min_height:.1f}".replace('.', ',')
                            max_str = f"{max_height:.3f}".replace('.', ',')
                            part_title = f"Часть {part_num} ({min_str} - {max_str})"
                            part_heading = doc.add_paragraph(part_title)
                            part_heading.runs[0].font.bold = True
                            part_heading.runs[0].font.size = Pt(11)
                            doc.add_paragraph()  # Отступ
                            
                            # Получаем все уникальные высоты для этой части
                            all_heights = set()
                            max_tolerance = 0
                            for belt_data in belts_data.values():
                                for item in belt_data:
                                    all_heights.add(round(item['height'], 1))
                                    max_tolerance = max(max_tolerance, item.get('tolerance', 0))
                            
                            sorted_heights = sorted(all_heights)
                            sorted_belts = sorted(belts_data.keys())
                            
                            # Создаем словарь для быстрого доступа
                            belt_height_deflection = {}
                            for belt_num, belt_data in belts_data.items():
                                for item in belt_data:
                                    height_rounded = round(item['height'], 1)
                                    belt_height_deflection[(belt_num, height_rounded)] = item.get('deflection', 0)
                            
                            # Создаем таблицу для этой части
                            deflection_table = doc.add_table(rows=len(sorted_heights) + 1, cols=len(sorted_belts) + 2)
                            deflection_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            # Заголовок
                            header_row = deflection_table.rows[0]
                            header_row.cells[0].text = 'Высота, м'
                            for j, belt_num in enumerate(sorted_belts, start=1):
                                header_row.cells[j].text = f'Пояс {belt_num}, мм'
                            header_row.cells[-1].text = 'Допуск, мм'
                            
                            for paragraph in header_row.cells[0].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                paragraph.runs[0].font.bold = True
                                paragraph.runs[0].font.size = Pt(9)
                            
                            for j in range(1, len(header_row.cells)):
                                for paragraph in header_row.cells[j].paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    paragraph.runs[0].font.bold = True
                                    paragraph.runs[0].font.size = Pt(9)
                            
                            # Данные
                            for i, height in enumerate(sorted_heights):
                                data_row = deflection_table.rows[i + 1]
                                data_row.cells[0].text = f"{height:.1f}"
                                
                                for j, belt_num in enumerate(sorted_belts, start=1):
                                    key = (belt_num, height)
                                    if key in belt_height_deflection:
                                        deflection = belt_height_deflection[key]
                                        data_row.cells[j].text = f"{deflection:+.1f}"
                                        
                                        # Выделяем превышения
                                        if abs(deflection) > max_tolerance:
                                            para = data_row.cells[j].paragraphs[0]
                                            para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                                            para.runs[0].font.bold = True
                                    else:
                                        data_row.cells[j].text = "—"
                                
                                data_row.cells[-1].text = f"±{max_tolerance:.1f}"
                                
                                # Форматирование
                                for cell in data_row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.size = Pt(8)
                            
                            format_table(
                                deflection_table,
                                header_rows=1,
                                font_size=8,
                                column_widths=[2.0] + [2.0] * len(sorted_belts) + [2.0]
                            )
                            
                            doc.add_paragraph()  # Отступ после таблицы части
                            
                            # Графики прямолинейности для этой части сразу после таблицы
                            heading_graphs = doc.add_paragraph('Графики прямолинейности')
                            heading_graphs.runs[0].font.bold = True
                            heading_graphs.runs[0].font.size = Pt(11)
                            doc.add_paragraph()  # Отступ
                            
                            # Получаем графики для этой части
                            part_figures = []
                            if hasattr(straightness_widget, 'get_part_figures_for_pdf'):
                                try:
                                    part_figures = straightness_widget.get_part_figures_for_pdf(part_num, group_size=2)
                                except Exception as exc:  # noqa: BLE001
                                    logger.warning(f"Не удалось получить графики для части {part_num}: {exc}")
                            
                            if not part_figures:
                                # Fallback: пробуем получить все графики и отфильтровать
                                try:
                                    if hasattr(straightness_widget, 'get_grouped_figures_for_pdf'):
                                        all_figures = straightness_widget.get_grouped_figures_for_pdf()
                                        # Фильтруем по поясам, принадлежащим этой части
                                        part_belts = set(sorted_belts)
                                        for belt_group, figure in all_figures:
                                            if any(belt in part_belts for belt in belt_group):
                                                part_figures.append((belt_group, figure))
                                except Exception as exc:  # noqa: BLE001
                                    logger.warning(f"Не удалось получить графики (fallback) для части {part_num}: {exc}")
                            
                            if part_figures:
                                for belt_group, figure in part_figures:
                                    width = 11.0 if len(belt_group) > 1 else 8.5
                                    original_size = self._prepare_matplotlib_figure(
                                        figure,
                                        width=width,
                                        height=5.8,
                                        pad=1.6,
                                        label_size=9,
                                        title_size=11
                                    )
                                    # Сохраняем в векторном формате (предпочтительно EMF)
                                    tmp_path_base = temp_dir / f'straightness_part_{part_num}_group_{global_figure_index}_docx'
                                    tmp_path = self._save_figure_vector_format(figure, tmp_path_base, preferred_format='emf')
                                    temp_files.append(str(tmp_path))
                                    
                                    # Вставляем график (python-docx поддерживает EMF, WMF и SVG)
                                    try:
                                        doc.add_picture(str(tmp_path), width=Inches(6))
                                    except Exception as e:
                                        # Если не удалось вставить векторный формат, используем высококачественный PNG как fallback
                                        logger.warning(f"Не удалось вставить векторный график прямолинейности: {e}. Используем PNG.")
                                        png_path = temp_dir / f'straightness_part_{part_num}_group_{global_figure_index}_docx.png'
                                        figure.savefig(str(png_path), dpi=600, bbox_inches='tight')
                                        temp_files.append(str(png_path))
                                        doc.add_picture(str(png_path), width=Inches(6))
                                    caption = doc.add_paragraph(
                                        f'Рис. {global_figure_index}. Отклонения от прямолинейности по поясам {", ".join(str(b) for b in belt_group)}'
                                    )
                                    caption.runs[0].italic = True
                                    caption.runs[0].font.size = Pt(9)
                                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    doc.add_paragraph()
                                    
                                    if original_size is not None:
                                        figure.set_size_inches(original_size)
                                    
                                    global_figure_index += 1
                            else:
                                no_graphs_par = doc.add_paragraph('Графики для этой части недоступны')
                                no_graphs_par.runs[0].italic = True
                                no_graphs_par.runs[0].font.size = Pt(10)
                            
                            doc.add_paragraph()  # Отступ после графиков части
                        
                except Exception as e:
                    print(f"Ошибка генерации таблицы прогибов в DOCX: {e}")
            else:
                no_data_par = doc.add_paragraph('Данные прямолинейности отсутствуют')
                no_data_par.runs[0].italic = True
                no_data_par.runs[0].font.size = Pt(10)
                doc.add_paragraph()
            
            # Заключение о прямолинейности после всех графиков
            if straightness_widget and hasattr(straightness_widget, 'get_all_belts_data'):
                try:
                    straightness_data_for_conclusion = straightness_widget.get_all_belts_data()
                    if straightness_data_for_conclusion:
                        # Находим максимальное значение прогиба
                        max_deflection = 0.0
                        max_height = 0.0
                        max_belt_num = None
                        max_tolerance = 0.0
                        
                        for part_num, part_info in straightness_data_for_conclusion.items():
                            belts_data = part_info.get('belts', {})
                            for belt_num, belt_data in belts_data.items():
                                for item in belt_data:
                                    deflection_abs = abs(item.get('deflection', 0))
                                    if deflection_abs > max_deflection:
                                        max_deflection = deflection_abs
                                        max_height = item.get('height', 0)
                                        max_belt_num = belt_num
                                        max_tolerance = item.get('tolerance', 0)
                        
                        if max_deflection > 0.0:
                            exceeds_text = "превышает" if max_deflection > max_tolerance else "не превышает"
                            height_str = f"+{max_height:.1f}" if max_height >= 0 else f"{max_height:.1f}"
                            
                            conclusion_heading = doc.add_paragraph()
                            conclusion_heading_run = conclusion_heading.add_run('Заключение:')
                            conclusion_heading_run.font.bold = True
                            conclusion_heading_run.font.size = Pt(10)
                            
                            conclusion_para1 = doc.add_paragraph()
                            conclusion_text1 = "Стрелы прогиба рассчитаны относительно базовой линии между нижним и верхним поясами. Значения сопоставлены с нормативом δ ≤ L / 750 (ГОСТ Р 71949-2025 «Конструкции опорные антенных сооружений объектов связи. Правила приемки работ и эксплуатации»)."
                            conclusion_para1.add_run(conclusion_text1).font.size = Pt(10)
                            
                            conclusion_para2 = doc.add_paragraph()
                            conclusion_text2 = f"Стрела прогиба поясов башни {exceeds_text} допустимые значения. Максимальное значение составляет {max_deflection:.1f} мм на отм. {height_str} пояса №{max_belt_num} при допустимом значении {max_tolerance:.1f} мм."
                            conclusion_para2.add_run(conclusion_text2).font.size = Pt(10)
                            
                            doc.add_paragraph()  # Отступ
                except Exception as e:
                    logger.warning(f"Ошибка формирования заключения о прямолинейности: {e}")

            # Сохраняем документ
            doc.save(output_path)
            
            # Удаляем временные файлы
            for temp_file in temp_files:
                try:
                    if Path(temp_file).exists():
                        Path(temp_file).unlink()
                except Exception as e:
                    print(f"Не удалось удалить временный файл {temp_file}: {e}")
            
            # Удаляем директорию, если она пуста
            try:
                if temp_dir.exists() and not any(temp_dir.iterdir()):
                    temp_dir.rmdir()
            except Exception as e:
                print(f"Не удалось удалить временную директорию {temp_dir}: {e}")
            
        except ImportError:
            raise ValueError("Для генерации DOCX отчетов требуется установить python-docx: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Ошибка генерации DOCX отчета: {str(e)}")

