"""
Генератор отчетов в Excel и PDF
"""

from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

from core.normatives import get_vertical_tolerance
from core.services.straightness_profiles import get_preferred_straightness_part_map


class ReportGenerator:
    """Генератор отчетов для анализа вертикальности мачт"""

    def __init__(self):
        self.timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    @staticmethod
    def _get_preferred_straightness_data(
        processed_data: dict | None,
        raw_data: pd.DataFrame | None = None,
        straightness_widget=None,
    ) -> dict[int, dict]:
        widget_data = {}
        if straightness_widget and hasattr(straightness_widget, 'get_all_belts_data'):
            try:
                widget_data = straightness_widget.get_all_belts_data()
            except Exception as exc:
                print(f"Не удалось получить данные прямолинейности из виджета: {exc}")

        return get_preferred_straightness_part_map(
            processed_data.get('straightness_profiles') if isinstance(processed_data, dict) else None,
            widget_data,
            points=raw_data,
            tower_parts_info=processed_data.get('tower_parts_info') if isinstance(processed_data, dict) else None,
        )

    @staticmethod
    def _summarize_straightness_data(straightness_data: dict[int, dict]) -> dict[str, Any]:
        summary = {
            'total': 0,
            'passed': 0,
            'failed': 0,
            'max_deflection_mm': 0.0,
            'max_tolerance_mm': 0.0,
            'max_height_m': 0.0,
            'max_belt_num': None,
            'max_part_num': None,
        }

        for part_num, part_info in (straightness_data or {}).items():
            belts_data = (part_info or {}).get('belts', {}) or {}
            for belt_num, belt_points in belts_data.items():
                for item in belt_points or []:
                    try:
                        height_m = float(item.get('height', 0.0) or 0.0)
                        deviation_mm = float(item.get('deflection', 0.0) or 0.0)
                        tolerance_mm = float(item.get('tolerance', 0.0) or 0.0)
                    except (TypeError, ValueError):
                        continue

                    abs_deviation = abs(deviation_mm)
                    summary['total'] += 1
                    if abs_deviation <= tolerance_mm + 1e-9:
                        summary['passed'] += 1
                    else:
                        summary['failed'] += 1

                    if abs_deviation > summary['max_deflection_mm']:
                        summary['max_deflection_mm'] = abs_deviation
                        summary['max_tolerance_mm'] = tolerance_mm
                        summary['max_height_m'] = height_m
                        summary['max_belt_num'] = int(belt_num)
                        summary['max_part_num'] = int(part_num)

        return summary

    @staticmethod
    def _resolve_straightness_at_height(
        straightness_data: dict[int, dict],
        target_height: float,
        tolerance: float = 0.25,
    ) -> float | None:
        matched_values: list[float] = []
        nearest_value: float | None = None
        nearest_delta: float | None = None

        for part_info in (straightness_data or {}).values():
            belts_data = (part_info or {}).get('belts', {}) or {}
            for belt_points in belts_data.values():
                for item in belt_points or []:
                    try:
                        height_m = float(item.get('height', 0.0) or 0.0)
                        deviation_mm = abs(float(item.get('deflection', 0.0) or 0.0))
                    except (TypeError, ValueError):
                        continue

                    delta = abs(height_m - target_height)
                    if delta <= tolerance:
                        matched_values.append(deviation_mm)
                    if nearest_delta is None or delta < nearest_delta:
                        nearest_delta = delta
                        nearest_value = deviation_mm

        if matched_values:
            return max(matched_values)
        return nearest_value

    @staticmethod
    def _iter_straightness_rows(straightness_data: dict[int, dict]):
        for part_num in sorted((straightness_data or {}).keys()):
            part_info = (straightness_data or {}).get(part_num) or {}
            belts_data = (part_info or {}).get('belts', {}) or {}
            for belt_num in sorted(belts_data.keys()):
                for item in belts_data.get(belt_num) or []:
                    if isinstance(item, dict):
                        yield int(part_num), int(belt_num), item

    @staticmethod
    def _prepare_matplotlib_figure(figure, width: float, height: float, *, pad: float = 1.2,
                                   label_size: int = 9, title_size: int = 11):
        """
        Настраивает matplotlib.figure перед экспортом и возвращает исходный размер
        для последующего восстановления.
        """
        if figure is None:
            return None

        original_size = figure.get_size_inches()
        figure.set_size_inches(width, height)

        for ax in figure.axes:
            ax.tick_params(labelsize=label_size)
            if ax.title and ax.title.get_text():
                ax.title.set_fontsize(title_size)
            if ax.xaxis and ax.xaxis.label:
                ax.xaxis.label.set_fontsize(label_size)
            if ax.yaxis and ax.yaxis.label:
                ax.yaxis.label.set_fontsize(label_size)
        figure.tight_layout(pad=pad)

        return original_size

    def generate_excel_report(self,
                              raw_data: pd.DataFrame,
                              processed_data: dict,
                              output_path: str,
                              angular_measurements: dict | None = None):
        """
        Генерирует отчет в формате Excel

        Args:
            raw_data: Исходные данные
            processed_data: Результаты расчетов
            output_path: Путь для сохранения
            angular_measurements: Данные угловых измерений по осям X и Y (опционально)
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
            from openpyxl.utils import get_column_letter

            wb = Workbook()

            # Лист 1: Результаты расчетов
            ws_results = wb.active
            ws_results.title = "Результаты"

            centers = processed_data['centers']
            straightness_data = self._get_preferred_straightness_data(processed_data, raw_data)
            straightness_summary = self._summarize_straightness_data(straightness_data)

            ws_results['A1'] = 'Результаты расчетов'
            ws_results['A1'].font = Font(size=14, bold=True)

            ws_results['A3'] = 'Центры поясов и отклонения'
            ws_results['A3'].font = Font(size=12, bold=True)

            # Заголовки
            headers = ['№ пояса', 'Высота (м)', 'X центра (м)', 'Y центра (м)',
                      'Отклонение от вертикали (мм)', 'Стрела прогиба (мм)',
                      'Точек в поясе']

            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            data_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            thin = Side(style='thin', color='FFBFBFBF')
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            for col_idx, header in enumerate(headers, start=1):
                cell = ws_results.cell(row=4, column=col_idx)
                cell.value = header
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
                cell.alignment = header_alignment
                cell.border = border

            ws_results.row_dimensions[4].height = 28

            # Данные
            for i, (idx, row) in enumerate(centers.iterrows(), start=5):
                ws_results.cell(row=i, column=1, value=i - 4)
                ws_results.cell(row=i, column=2, value=round(row['z'], 3))
                ws_results.cell(row=i, column=3, value=round(row['x'], 6))
                ws_results.cell(row=i, column=4, value=round(row['y'], 6))
                ws_results.cell(row=i, column=5, value=round(row['deviation'] * 1000, 2))

                straightness_mm = self._resolve_straightness_at_height(straightness_data, float(row['z']))
                if straightness_mm is None and 'straightness_deviation' in row:
                    straightness_mm = abs(float(row['straightness_deviation']) * 1000.0)
                if straightness_mm is not None:
                    ws_results.cell(row=i, column=6, value=round(straightness_mm, 2))

                if 'points_count' in row:
                    ws_results.cell(row=i, column=7, value=int(row['points_count']))

                for col_idx in range(1, len(headers) + 1):
                    data_cell = ws_results.cell(row=i, column=col_idx)
                    data_cell.alignment = data_alignment
                    data_cell.border = border
                ws_results.row_dimensions[i].height = 20

            for col_idx, width in enumerate([10, 14, 16, 16, 20, 18, 14], start=1):
                ws_results.column_dimensions[get_column_letter(col_idx)].width = width

            ws_results.freeze_panes = 'A5'

            # Лист 2: Нормативы и выводы
            ws3 = wb.create_sheet("Нормативы")

            ws3['A1'] = 'Проверка соответствия нормативам'
            ws3['A1'].font = Font(size=14, bold=True)

            ws3['A3'] = 'Нормативная база:'
            ws3['A4'] = 'СП 70.13330.2012: Отклонение от вертикали d ≤ 0.001 × h'
            ws3['A5'] = 'ГОСТ Р 71949-2025 Конструкции опорные антенных сооружений объектов связи: Стрела прогиба δ ≤ L / 750'
            for cell in ['A3', 'A4', 'A5']:
                ws3[cell].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

            # Статистика
            from core.normatives import NormativeChecker
            checker = NormativeChecker()

            vertical_check = checker.check_vertical_deviations(
                centers['deviation'].tolist(),
                centers['z'].tolist()
            )

            ws3['A7'] = 'Результаты проверки вертикальности:'
            ws3['A8'] = f"Всего поясов: {vertical_check['total']}"
            ws3['A9'] = f"✓ В норме: {vertical_check['passed']}"
            ws3['A10'] = f"✗ Превышение: {vertical_check['failed']}"

            ws3['A9'].font = Font(color='008000')
            ws3['A10'].font = Font(color='FF0000')

            if vertical_check['non_compliant']:
                ws3['A12'] = 'Пояса с превышением допуска:'
                row_idx = 13
                for item in vertical_check['non_compliant']:
                    excess_mm = item['excess'] * 1000
                    ws3[f'A{row_idx}'] = f"Пояс {item['index']+1} (h={item['height']:.1f}м): превышение {excess_mm:.1f} мм"
                    ws3[f'A{row_idx}'].font = Font(color='FF0000')
                    ws3[f'A{row_idx}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                    row_idx += 1

            # Лист 3: Журнал угловых измерений
            if not vertical_check['non_compliant']:
                row_idx = 12

            if straightness_summary['total'] > 0:
                row_idx += 1
                ws3[f'A{row_idx}'] = '\u0420\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u044b \u043f\u0440\u043e\u0432\u0435\u0440\u043a\u0438 \u043f\u0440\u044f\u043c\u043e\u043b\u0438\u043d\u0435\u0439\u043d\u043e\u0441\u0442\u0438 \u043f\u043e\u044f\u0441\u043e\u0432:'
                ws3[f'A{row_idx + 1}'] = f"\u0412\u0441\u0435\u0433\u043e \u0442\u043e\u0447\u0435\u043a: {straightness_summary['total']}"
                ws3[f'A{row_idx + 2}'] = f"\u2713 \u0412 \u043d\u043e\u0440\u043c\u0435: {straightness_summary['passed']}"
                ws3[f'A{row_idx + 3}'] = f"\u2717 \u041f\u0440\u0435\u0432\u044b\u0448\u0435\u043d\u0438\u0435: {straightness_summary['failed']}"
                ws3[f'A{row_idx + 2}'].font = Font(color='008000')
                ws3[f'A{row_idx + 3}'].font = Font(color='FF0000')

                max_location = (
                    f"\u043f\u043e\u044f\u0441 {straightness_summary['max_belt_num']}, "
                    f"\u0432\u044b\u0441\u043e\u0442\u0430 {straightness_summary['max_height_m']:.3f} \u043c"
                )
                if len(straightness_data) > 1 and straightness_summary['max_part_num'] is not None:
                    max_location = f"\u0447\u0430\u0441\u0442\u044c {straightness_summary['max_part_num']}, {max_location}"

                ws3[f'A{row_idx + 5}'] = (
                    '\u041c\u0430\u043a\u0441\u0438\u043c\u0430\u043b\u044c\u043d\u0430\u044f \u0441\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430: '
                    f"{straightness_summary['max_deflection_mm']:.1f} \u043c\u043c "
                    f'\u043f\u0440\u0438 \u0434\u043e\u043f\u0443\u0441\u043a\u0435 {straightness_summary["max_tolerance_mm"]:.1f} \u043c\u043c '
                    f'({max_location})'
                )
                ws3[f'A{row_idx + 5}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

                ws_straight = wb.create_sheet("\u041f\u0440\u044f\u043c\u043e\u043b\u0438\u043d\u0435\u0439\u043d\u043e\u0441\u0442\u044c")
                ws_straight['A1'] = '\u0420\u0430\u0441\u0447\u0435\u0442 \u0441\u0442\u0440\u0435\u043b\u044b \u043f\u0440\u043e\u0433\u0438\u0431\u0430 \u043f\u043e\u044f\u0441\u043e\u0432'
                ws_straight['A1'].font = Font(size=14, bold=True)

                straight_headers = [
                    '\u0427\u0430\u0441\u0442\u044c',
                    '\u2116 \u043f\u043e\u044f\u0441\u0430',
                    '\u0412\u044b\u0441\u043e\u0442\u0430 (\u043c)',
                    '\u0421\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430 (\u043c\u043c)',
                    '\u0414\u043e\u043f\u0443\u0441\u043a (\u043c\u043c)',
                    '\u0421\u043e\u0441\u0442\u043e\u044f\u043d\u0438\u0435',
                ]

                for col_idx, header in enumerate(straight_headers, start=1):
                    cell = ws_straight.cell(row=3, column=col_idx)
                    cell.value = header
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
                    cell.alignment = header_alignment
                    cell.border = border

                straight_row_idx = 4
                for part_num, belt_num, item in self._iter_straightness_rows(straightness_data):
                    height_m = float(item.get('height', 0.0) or 0.0)
                    deflection_mm = float(item.get('deflection', 0.0) or 0.0)
                    tolerance_mm = float(item.get('tolerance', 0.0) or 0.0)
                    is_within_limit = abs(deflection_mm) <= tolerance_mm + 1e-9
                    status = '\u0421\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0443\u0435\u0442' if is_within_limit else '\u041f\u0440\u0435\u0432\u044b\u0448\u0435\u043d\u0438\u0435'

                    values = [
                        part_num,
                        belt_num,
                        round(height_m, 3),
                        round(deflection_mm, 2),
                        round(tolerance_mm, 2),
                        status,
                    ]
                    for col_idx, value in enumerate(values, start=1):
                        cell = ws_straight.cell(row=straight_row_idx, column=col_idx, value=value)
                        cell.alignment = data_alignment
                        cell.border = border

                    ws_straight.cell(row=straight_row_idx, column=6).font = Font(
                        color='008000' if is_within_limit else 'FF0000',
                        bold=not is_within_limit,
                    )
                    straight_row_idx += 1

                for col_idx, width in enumerate([10, 12, 14, 20, 16, 18], start=1):
                    ws_straight.column_dimensions[get_column_letter(col_idx)].width = width

                ws_straight.freeze_panes = 'A4'

            if angular_measurements and (angular_measurements.get('x') or angular_measurements.get('y')):
                ws_angular = wb.create_sheet("Журнал угловых измерений")

                headers_angular = ['№', 'Секция', 'H, м', 'Пояс', 'KL', 'KR', 'KL–KR (″)', 'βизм', 'Bизм', 'Δβ', 'Δb, мм']
                # Оптимальные ширины столбцов (в единицах Excel, примерно соответствуют см)
                angular_col_widths = [5, 9, 8, 12, 10, 10, 12, 10, 10, 10, 10]

                def append_angular_table_sheet(axis_label: str, rows: list, start_row: int):
                    """Добавляет таблицу угловых измерений для одной оси на лист."""
                    current_row = start_row

                    # Заголовок оси
                    ws_angular[f'A{current_row}'] = f'Ось {axis_label}'
                    ws_angular[f'A{current_row}'].font = Font(size=12, bold=True)
                    current_row += 2

                    if not rows:
                        ws_angular[f'A{current_row}'] = 'Данные отсутствуют'
                        ws_angular[f'A{current_row}'].font = Font(italic=True)
                        return current_row + 2

                    # Заголовки таблицы
                    for col_idx, header in enumerate(headers_angular, start=1):
                        cell = ws_angular.cell(row=current_row, column=col_idx)
                        # Для столбца "Пояс" разбиваем заголовок на две строки
                        if col_idx == 4:
                            cell.value = 'По\nяс'
                        else:
                            cell.value = header
                        cell.font = Font(bold=True, size=9)
                        cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        cell.border = border

                    ws_angular.row_dimensions[current_row].height = 40

                    # Ширина столбцов
                    for col_idx, width in enumerate(angular_col_widths, start=1):
                        ws_angular.column_dimensions[get_column_letter(col_idx)].width = width

                    current_row += 1

                    # Данные
                    for idx, row in enumerate(rows, start=1):
                        height = row.get('height')
                        height_str = f"{float(height):.3f}" if height is not None else '—'
                        belt_value = row.get('belt', '—')
                        belt_str = str(belt_value) if belt_value is not None else '—'

                        values = [
                            str(idx),
                            str(row.get('section_label', '—')),
                            height_str,
                            belt_str,
                            row.get('kl_str', '—'),
                            row.get('kr_str', '—'),
                            row.get('diff_str', '—'),
                            row.get('beta_str', '—'),
                            row.get('center_str', '—'),
                            row.get('delta_str', '—'),
                            row.get('delta_mm_str', '—'),
                        ]

                        for col_idx, value in enumerate(values, start=1):
                            cell = ws_angular.cell(row=current_row, column=col_idx)
                            cell.value = value
                            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                            cell.border = border
                            # Для столбца "Пояс" включаем перенос текста
                            if col_idx == 4:
                                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

                        ws_angular.row_dimensions[current_row].height = 20
                        current_row += 1

                    return current_row + 2

                # Добавляем таблицы для осей X и Y
                rows_x = angular_measurements.get('x', [])
                rows_y = angular_measurements.get('y', [])

                start_row = 1
                if rows_x:
                    start_row = append_angular_table_sheet('X', rows_x, start_row)
                if rows_y:
                    start_row = append_angular_table_sheet('Y', rows_y, start_row)

                # Замораживаем первую строку с заголовками
                ws_angular.freeze_panes = 'A1'

            # Автоширина колонок (кроме листа с угловыми измерениями, там ширина задана явно)
            sheets_for_auto_width = [ws_results, ws3]
            for ws in sheets_for_auto_width:
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except (TypeError, ValueError):
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width

            # Сохраняем
            wb.save(output_path)

        except Exception as e:
            raise ValueError(f"Ошибка генерации Excel отчета: {e!s}")

    def generate_pdf_report(self,
                           raw_data: pd.DataFrame,
                           processed_data: dict,
                           output_path: str,
                           vertical_plot_widget=None,
                           straightness_plot_widget=None):
        """
        Генерирует отчет в формате PDF

        Args:
            raw_data: Исходные данные
            processed_data: Результаты расчетов
            output_path: Путь для сохранения
            vertical_plot_widget: Виджет графика вертикальности
            straightness_plot_widget: Виджет графика прямолинейности
        """
        try:
            import tempfile

            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

            # Создаем PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()

            # Титульный лист
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#2C3E50'),
                spaceAfter=12,
                spaceBefore=60,
                alignment=1  # CENTER
            )

            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=14,
                textColor=colors.HexColor('#34495E'),
                spaceAfter=30,
                alignment=1
            )

            story.append(Paragraph('<b>ОТЧЕТ</b>', title_style))
            story.append(Paragraph('ПО ГЕОДЕЗИЧЕСКОМУ КОНТРОЛЮ', title_style))
            story.append(Paragraph('АНТЕННО-МАЧТОВОГО СООРУЖЕНИЯ', title_style))
            story.append(Spacer(1, 40))
            story.append(Paragraph(f'Дата обследования: {self.timestamp}', subtitle_style))
            story.append(Spacer(1, 20))

            # Секция: Нормативная база
            section_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#16A085'),
                spaceBefore=20,
                spaceAfter=10,
                leftIndent=0
            )

            story.append(PageBreak())
            story.append(Paragraph('<b>1. НОРМАТИВНАЯ БАЗА</b>', section_style))
            story.append(Paragraph('<b>СП 70.13330.2012</b> "Несущие и ограждающие конструкции":', styles['Normal']))
            story.append(Paragraph('&nbsp;&nbsp;&nbsp;• Отклонение от вертикали: <b>d ≤ 0.001 × h</b>', styles['Normal']))
            story.append(Spacer(1, 5))
            story.append(Paragraph('<b>ГОСТ Р 71949-2025 Конструкции опорные антенных сооружений объектов связи:</b>', styles['Normal']))
            story.append(Paragraph('&nbsp;&nbsp;&nbsp;• Стрела прогиба: <b>δ ≤ L / 750</b>', styles['Normal']))
            story.append(Spacer(1, 20))

            story.append(PageBreak())

            # Секция: Результаты расчетов
            story.append(Paragraph('2. Результаты расчетов', styles['Heading2']))

            centers = processed_data['centers']
            story.append(Paragraph(f'Количество обнаруженных поясов: {len(centers)}', styles['Normal']))
            story.append(Spacer(1, 10))

            # Таблица результатов (улучшенный стиль)
            data = [['№\nпояса', 'Высота,\nм', 'Отклонение\nот вертикали,\nмм', 'Допуск,\nмм', 'Соответствие\nнормативу']]

            straightness_data = self._get_preferred_straightness_data(
                processed_data,
                raw_data,
                straightness_plot_widget,
            )
            straightness_summary = self._summarize_straightness_data(straightness_data)

            from core.normatives import get_vertical_tolerance

            for i, (idx, row) in enumerate(centers.iterrows(), start=1):
                dev_mm = row['deviation'] * 1000
                tolerance_mm = get_vertical_tolerance(row['z']) * 1000
                status = '✓ Соответствует' if abs(dev_mm) <= tolerance_mm else '✗ Превышение'

                data.append([
                    str(i),
                    f"{row['z']:.2f}",
                    f"{dev_mm:.2f}",
                    f"{tolerance_mm:.2f}",
                    status
                ])

            table = Table(data, colWidths=[1.5*cm, 2*cm, 3*cm, 2.5*cm, 4*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2C3E50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ECF0F1')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#95A5A6')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')])
            ]))

            story.append(table)
            story.append(PageBreak())

            # Секция: Графики
            story.append(Paragraph('3. Графический анализ', styles['Heading2']))

            # Сохраняем графики во временные файлы
            if vertical_plot_widget and hasattr(vertical_plot_widget, 'figure'):
                fig = vertical_plot_widget.figure
                original_size = self._prepare_matplotlib_figure(fig, 7.5, 6.0, pad=1.4)
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    fig.savefig(tmp.name, dpi=220, bbox_inches='tight')
                    img = Image(tmp.name, width=15*cm, height=12*cm)
                    story.append(img)
                    Path(tmp.name).unlink()
                if original_size is not None:
                    fig.set_size_inches(original_size)

            story.append(Spacer(1, 10))

            grouped_straightness_figures = []
            if (
                straightness_plot_widget
                and len(straightness_data) > 1
                and hasattr(straightness_plot_widget, 'get_part_figures_for_pdf')
            ):
                for part_num in sorted(straightness_data.keys()):
                    try:
                        for belt_group, fig in straightness_plot_widget.get_part_figures_for_pdf(part_num):
                            grouped_straightness_figures.append((part_num, belt_group, fig))
                    except Exception as exc:
                        print(
                            f"РќРµ СѓРґР°Р»РѕСЃСЊ РїРѕР»СѓС‡РёС‚СЊ РіСЂСѓРїРїРѕРІС‹Рµ "
                            f"РіСЂР°С„РёРєРё РїСЂСЏРјРѕР»РёРЅРµР№РЅРѕСЃС‚Рё РґР»СЏ С‡Р°СЃС‚Рё {part_num}: {exc}"
                        )
            if not grouped_straightness_figures and straightness_plot_widget and hasattr(straightness_plot_widget, 'get_grouped_figures_for_pdf'):
                try:
                    grouped_straightness_figures = [
                        (None, belt_group, fig)
                        for belt_group, fig in straightness_plot_widget.get_grouped_figures_for_pdf()
                    ]
                except Exception as exc:
                    print(f"Не удалось получить сгруппированные графики прямолинейности: {exc}")

            if grouped_straightness_figures:
                figure_index = 2
                for part_num, belt_group, fig in grouped_straightness_figures:
                    width = 9.0 if len(belt_group) == 1 else 10.5
                    original_size = self._prepare_matplotlib_figure(fig, width, 5.8, pad=1.6)
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        fig.savefig(tmp.name, dpi=220, bbox_inches='tight')
                        img = Image(tmp.name, width=15*cm, height=11*cm)
                        story.append(img)
                        Path(tmp.name).unlink()
                    if original_size is not None:
                        fig.set_size_inches(original_size)

                    belts_caption = ', '.join(str(b) for b in belt_group)
                    if part_num is not None and len(straightness_data) > 1:
                        belts_caption = f'\u0447\u0430\u0441\u0442\u0438 {part_num}, {belts_caption}'
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(
                        f'<i>Рис. {figure_index}. Отклонения от прямолинейности по поясам {belts_caption}</i>',
                        ParagraphStyle('Caption', parent=styles['Normal'], fontSize=9, alignment=1) # TA_CENTER
                    ))
                    story.append(Spacer(1, 14))
                    figure_index += 1
            elif straightness_plot_widget and hasattr(straightness_plot_widget, 'figure'):
                fig = straightness_plot_widget.figure
                original_size = self._prepare_matplotlib_figure(fig, 9.0, 6.5, pad=1.6)
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    fig.savefig(tmp.name, dpi=220, bbox_inches='tight')
                    img = Image(tmp.name, width=15*cm, height=12*cm)
                    story.append(img)
                    Path(tmp.name).unlink()
                if original_size is not None:
                    fig.set_size_inches(original_size)

            story.append(PageBreak())

            # Секция: Выводы
            story.append(Paragraph('4. Выводы', styles['Heading2']))

            from core.normatives import NormativeChecker
            checker = NormativeChecker()

            vertical_check = checker.check_vertical_deviations(
                centers['deviation'].tolist(),
                centers['z'].tolist()
            )

            story.append(Paragraph(f"Проверено поясов: {vertical_check['total']}", styles['Normal']))
            story.append(Paragraph(f"✓ Соответствуют нормативам: {vertical_check['passed']}", styles['Normal']))
            story.append(Paragraph(f"✗ Превышение допуска: {vertical_check['failed']}", styles['Normal']))

            if vertical_check['failed'] == 0:
                story.append(Spacer(1, 10))
                story.append(Paragraph('<b>Заключение: Мачта соответствует нормативным требованиям по вертикальности.</b>',
                                      styles['Normal']))
            else:
                story.append(Spacer(1, 10))
                story.append(Paragraph('<b>Заключение: Обнаружены превышения допустимых отклонений. Требуется корректировка.</b>',
                                      styles['Normal']))

            # Генерируем PDF
            if straightness_summary['total'] > 0:
                story.append(Spacer(1, 12))
                story.append(Paragraph('\u041f\u0440\u043e\u0432\u0435\u0440\u043a\u0430 \u043f\u0440\u044f\u043c\u043e\u043b\u0438\u043d\u0435\u0439\u043d\u043e\u0441\u0442\u0438 \u043f\u043e\u044f\u0441\u043e\u0432:', styles['Normal']))
                story.append(Paragraph(f"\u041f\u0440\u043e\u0432\u0435\u0440\u0435\u043d\u043e \u0442\u043e\u0447\u0435\u043a: {straightness_summary['total']}", styles['Normal']))
                story.append(Paragraph(f"\u2713 \u0412 \u043d\u043e\u0440\u043c\u0435: {straightness_summary['passed']}", styles['Normal']))
                story.append(Paragraph(f"\u2717 \u041f\u0440\u0435\u0432\u044b\u0448\u0435\u043d\u0438\u0435: {straightness_summary['failed']}", styles['Normal']))

                max_location = (
                    f"\u043f\u043e\u044f\u0441 {straightness_summary['max_belt_num']}, "
                    f"\u0432\u044b\u0441\u043e\u0442\u0430 {straightness_summary['max_height_m']:.3f} \u043c"
                )
                if len(straightness_data) > 1 and straightness_summary['max_part_num'] is not None:
                    max_location = f"\u0447\u0430\u0441\u0442\u044c {straightness_summary['max_part_num']}, {max_location}"

                if straightness_summary['failed'] == 0:
                    story.append(Paragraph(
                        '<b>\u0417\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435 \u043f\u043e \u043f\u0440\u044f\u043c\u043e\u043b\u0438\u043d\u0435\u0439\u043d\u043e\u0441\u0442\u0438: '
                        '\u0441\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430 \u043f\u043e\u044f\u0441\u043e\u0432 \u043d\u0435 \u043f\u0440\u0435\u0432\u044b\u0448\u0430\u0435\u0442 \u0434\u043e\u043f\u0443\u0441\u0442\u0438\u043c\u044b\u0445 '
                        '\u0437\u043d\u0430\u0447\u0435\u043d\u0438\u0439.</b>',
                        styles['Normal'],
                    ))
                else:
                    story.append(Paragraph(
                        '<b>\u0417\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435 \u043f\u043e \u043f\u0440\u044f\u043c\u043e\u043b\u0438\u043d\u0435\u0439\u043d\u043e\u0441\u0442\u0438: '
                        f"\u043c\u0430\u043a\u0441\u0438\u043c\u0430\u043b\u044c\u043d\u0430\u044f \u0441\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430 \u0441\u043e\u0441\u0442\u0430\u0432\u043b\u044f\u0435\u0442 "
                        f"{straightness_summary['max_deflection_mm']:.1f} \u043c\u043c \u043f\u0440\u0438 \u0434\u043e\u043f\u0443\u0441\u043a\u0435 "
                        f"{straightness_summary['max_tolerance_mm']:.1f} \u043c\u043c ({max_location}).</b>",
                        styles['Normal'],
                    ))

            doc.build(story)

        except Exception as e:
            raise ValueError(f"Ошибка генерации PDF отчета: {e!s}")

    def generate_docx_report(self,
                            raw_data: pd.DataFrame,
                            processed_data: dict,
                            output_path: str,
                            object_info: dict = None,
                            verticality_widget=None,
                            straightness_widget=None):
        """
        Генерирует отчет в формате DOCX (MS Word), идентичный PDF

        Args:
            raw_data: Исходные данные
            processed_data: Результаты расчетов
            output_path: Путь для сохранения
            object_info: Информация об объекте (название, адрес, и т.д.)
            verticality_widget: Виджет графика вертикальности
            straightness_widget: Виджет графика прямолинейности
        """
        try:
            import tempfile
            from pathlib import Path

            from docx import Document
            from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Inches, Pt, RGBColor

            # Создаем директорию для временных файлов
            output_dir = Path(output_path).parent
            temp_dir = output_dir / 'temp_charts'
            temp_dir.mkdir(exist_ok=True)
            temp_files = []

            # Создаем документ
            doc = Document()

            # Настройка полей страницы: левое 2,0 см, правое 0,75 см
            section = doc.sections[0]
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(0.75)

            def format_table(table, *, header_rows: int = 1, font_size: int = 9, column_widths=None):
                """Унифицированное форматирование таблиц DOCX для предотвращения наложений."""
                if table is None:
                    return

                table.autofit = True
                table.allow_autofit = True

                if column_widths:
                    for col_idx, width_cm in enumerate(column_widths):
                        if col_idx >= len(table.columns):
                            break
                        for cell in table.columns[col_idx].cells:
                            cell.width = Cm(width_cm)

                for row_idx, row in enumerate(table.rows):
                    row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.15
                            if paragraph.alignment is None:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                if run.font.size is None:
                                    run.font.size = Pt(font_size)
                        if row_idx < header_rows and cell.paragraphs:
                            for run in cell.paragraphs[0].runs:
                                run.font.bold = True

            # Настройка стилей документа
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'  # Используем Arial вместо Times New Roman
            font.size = Pt(10)

            # === ТИТУЛЬНАЯ СТРАНИЦА (Приложение Л) ===
            # Приложение Л (курсивом)
            title = doc.add_paragraph('Приложение Л', style='Title')
            title.runs[0].italic = True
            title.runs[0].font.size = Pt(12)

            # Подзаголовок
            subtitle = doc.add_paragraph('Протокол геодезических измерений сооружения', style='Body Text')
            subtitle.runs[0].font.size = Pt(12)

            doc.add_paragraph()  # Отступ

            centers = processed_data['centers']
            max_height = float(centers['z'].max()) if len(centers) > 0 else 0.0
            straightness_data = self._get_preferred_straightness_data(
                processed_data,
                raw_data,
                straightness_widget,
            )
            straightness_summary = self._summarize_straightness_data(straightness_data)

            # === ТАБЛИЦА ОТКЛОНЕНИЙ СТВОЛА ОТ ВЕРТИКАЛИ ===
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run('ТАБЛИЦА ОТКЛОНЕНИЙ СТВОЛА ОТ ВЕРТИКАЛИ')
            run.italic = True
            run.font.size = Pt(12)

            doc.add_paragraph()

            # Таблица с информацией об объекте
            info_table = doc.add_table(rows=2, cols=6)
            info_table.style = 'Table Grid'
            info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Первая строка
            row1 = info_table.rows[0]
            row1.cells[0].text = 'Тип опоры'
            row1.cells[1].text = object_info.get('project_name', 'Башня') if object_info else 'Башня'
            row1.cells[2].text = 'Высота опоры'
            row1.cells[3].text = f'{max_height:.1f} м'
            row1.cells[4].text = 'Инструмент'
            row1.cells[5].text = 'Тахеометр'

            # Вторая строка
            row2 = info_table.rows[1]
            row2.cells[0].text = 'Проект'
            row2.cells[1].text = 'ГСПИ'
            row2.cells[2].text = 'Город (поселок)'
            row2.cells[3].text = object_info.get('location', '') if object_info else ''
            row2.cells[4].text = 'Дата'
            row2.cells[5].text = self.timestamp.split()[0] if ' ' in self.timestamp else self.timestamp

            # Форматирование таблицы
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            format_table(
                info_table,
                header_rows=0,
                font_size=10,
                column_widths=[3.0, 4.0, 3.0, 3.2, 3.0, 4.0]
            )

            doc.add_paragraph()

            # Таблица с результатами отклонений (горизонтальная)
            results_table = doc.add_table(rows=2, cols=min(len(centers) + 1, 11))
            results_table.style = 'Table Grid'
            results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Заголовки (первые 10 поясов)
            header_row = results_table.rows[0]
            header_row.cells[0].text = '№ пояса'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                header_row.cells[i].text = str(i)

            # Отметки высот
            height_row = results_table.rows[1]
            height_row.cells[0].text = 'Отметка, м'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                height_row.cells[i].text = f"{row['z']:.3f}"

            # Отклонения от вертикали
            deviation_row = results_table.add_row()
            deviation_row.cells[0].text = 'Смещение центра сечения пояса от 0,001H вертикали, мм'

            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                dev_mm = row['deviation'] * 1000
                tolerance_mm = get_vertical_tolerance(row['z']) * 1000
                deviation_row.cells[i].text = f"{dev_mm:.1f}"

                # Красный цвет для превышений
                if abs(dev_mm) > tolerance_mm:
                    for paragraph in deviation_row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.bold = True

            results_columns = len(results_table.columns)
            format_table(
                results_table,
                header_rows=1,
                font_size=9,
                column_widths=[2.2] + [2.0] * (max(results_columns - 1, 0))
            )

            doc.add_paragraph()

            # Подпись составителя
            signature = doc.add_paragraph('Таблицу составил        ', style='Body Text')
            signature.add_run(object_info.get('executor', '____________') if object_info else '____________')
            signature.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            # === ПРОТОКОЛ ИЗМЕРЕНИЙ ВЕРТИКАЛЬНОСТИ СТВОЛА ОПОРЫ ===
            heading2 = doc.add_paragraph()
            heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = heading2.add_run('ПРОТОКОЛ ИЗМЕРЕНИЙ ВЕРТИКАЛЬНОСТИ СТВОЛА ОПОРЫ')
            run2.italic = True
            run2.font.size = Pt(12)

            doc.add_paragraph()

            # Таблица с параметрами измерений
            params_table = doc.add_table(rows=3, cols=6)
            params_table.style = 'Table Grid'
            params_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            row1 = params_table.rows[0]
            row1.cells[0].text = 'Тип опоры'
            row1.cells[1].text = object_info.get('project_name', 'Башня') if object_info else 'Башня'
            row1.cells[2].text = 'Высота опоры'
            row1.cells[3].text = f'{max_height:.1f} м'
            row1.cells[4].text = 'Инструмент'
            row1.cells[5].text = 'Тахеометр'

            row2 = params_table.rows[1]
            row2.cells[0].text = 'Облачность'
            row2.cells[1].text = ''
            row2.cells[2].text = 'Ветер'
            row2.cells[3].text = ''
            row2.cells[4].text = 'Изображения'
            row2.cells[5].text = 'Без искажений'

            row3 = params_table.rows[2]
            row3.cells[0].text = 'Наблюдатель'
            row3.cells[1].text = object_info.get('executor', '') if object_info else ''
            row3.cells[2].text = 'Дата'
            row3.cells[3].text = self.timestamp.split()[0] if ' ' in self.timestamp else self.timestamp
            row3.cells[4].text = ''
            row3.cells[5].text = ''

            format_table(
                params_table,
                header_rows=0,
                font_size=10,
                column_widths=[3.0, 3.8, 3.0, 3.2, 3.0, 3.0]
            )

            doc.add_paragraph()

            # Таблица с отклонениями по высотам (первые 10 поясов)
            heights_table = doc.add_table(rows=4, cols=min(len(centers) + 1, 11))
            heights_table.style = 'Table Grid'
            heights_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Строка 1: Высота сечения
            row1 = heights_table.rows[0]
            row1.cells[0].text = 'Высота сечения, м'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                row1.cells[i].text = f"{row['z']:.3f}"

            # Строка 2: Отклонение по оси X
            row2 = heights_table.rows[1]
            row2.cells[0].text = 'Отклонение от вертикали по оси Х, мм'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                x_dev = (row['x'] - centers.iloc[0]['x']) * 1000 if 'x' in row else row['deviation'] * 1000
                row2.cells[i].text = f"{x_dev:.1f}"

            # Строка 3: Отклонение по оси Y
            row3 = heights_table.rows[2]
            row3.cells[0].text = 'Отклонение от вертикали по оси Y, мм'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                y_dev = (row['y'] - centers.iloc[0]['y']) * 1000 if 'y' in row else 0.0
                row3.cells[i].text = f"{y_dev:.1f}"

            # Строка 4: Результирующее отклонение
            row4 = heights_table.rows[3]
            row4.cells[0].text = 'Результирующее отклонение, мм'
            for i, (idx, row) in enumerate(centers.head(10).iterrows(), start=1):
                dev_mm = row['deviation'] * 1000
                row4.cells[i].text = f"{dev_mm:.1f}"

            format_table(
                heights_table,
                header_rows=1,
                font_size=9,
                column_widths=[3.0] + [2.2] * (max(len(heights_table.columns) - 1, 0))
            )

            doc.add_paragraph()

            # === ЗАКЛЮЧЕНИЕ ===
            conclusion_heading = doc.add_paragraph('Заключение:', style='Body Text')
            conclusion_heading.runs[0].font.bold = True

            # Подсчет статистики
            compliant_count = sum(1 for _, row in centers.iterrows()
                                if abs(row['deviation'] * 1000) <= get_vertical_tolerance(row['z']) * 1000)
            non_compliant_count = len(centers) - compliant_count

            # Вывод 1: Соответствие нормативам
            conclusion1 = doc.add_paragraph(style='List Paragraph')
            if non_compliant_count == 0:
                conclusion1.add_run('Отклонения ствола от вертикали ').font.size = Pt(12)
                run = conclusion1.add_run('не превышают')
                run.font.size = Pt(12)
                run.font.bold = True
                conclusion1.add_run(' допусков СП 70.13330.2012 «Несущие и ограждающие конструкции».').font.size = Pt(12)
            else:
                conclusion1.add_run('Отклонения ствола от вертикали ').font.size = Pt(12)
                run = conclusion1.add_run('превышают')
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                conclusion1.add_run(f' допусков СП 70.13330.2012 «Несущие и ограждающие конструкции» на {non_compliant_count} поясах.').font.size = Pt(12)

            # Вывод 2: Эксплуатация
            conclusion2 = doc.add_paragraph(style='List Paragraph')
            if non_compliant_count == 0:
                conclusion2.add_run('Зафиксированные отклонения ствола от вертикали ').font.size = Pt(12)
                run = conclusion2.add_run('не препятствуют')
                run.font.size = Pt(12)
                run.font.bold = True
                conclusion2.add_run(' нормальной эксплуатации опоры.').font.size = Pt(12)
            else:
                conclusion2.add_run('Зафиксированные отклонения ствола от вертикали ').font.size = Pt(12)
                run = conclusion2.add_run('препятствуют')
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                conclusion2.add_run(' нормальной эксплуатации опоры.').font.size = Pt(12)

            doc.add_paragraph()

            # Подписи
            signature_para = doc.add_paragraph('Измерения выполнил              ', style='Body Text')
            signature_para.add_run(object_info.get('executor', '____________') if object_info else '____________')
            signature_para.add_run('        Вычисления проверил        ____________')

            # === РАСЧЕТ СТРЕЛЫ ПРОГИБА (если есть данные) ===
            if straightness_summary['total'] > 0:
                doc.add_page_break()

                heading3 = doc.add_paragraph()
                heading3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run3 = heading3.add_run('РАСЧЕТ СТРЕЛЫ ПРОГИБА ПОЯСА СТВОЛА (ПРЯМОЛИНЕЙНОСТЬ СТВОЛА)')
                run3.italic = True
                run3.font.size = Pt(12)

                doc.add_paragraph()

                # Таблица с параметрами
                straight_params = doc.add_table(rows=3, cols=6)
                straight_params.style = 'Table Grid'
                straight_params.alignment = WD_TABLE_ALIGNMENT.CENTER

                row1 = straight_params.rows[0]
                row1.cells[0].text = 'Тип опоры'
                row1.cells[1].text = object_info.get('project_name', 'Башня') if object_info else 'Башня'
                row1.cells[2].text = 'Высота опоры'
                row1.cells[3].text = f'{max_height:.1f} м'
                row1.cells[4].text = 'Инструмент'
                row1.cells[5].text = 'Тахеометр'

                row2 = straight_params.rows[1]
                row2.cells[0].text = 'Облачность'
                row2.cells[1].text = ''
                row2.cells[2].text = 'Ветер'
                row2.cells[3].text = ''
                row2.cells[4].text = 'Изображения'
                row2.cells[5].text = 'Без искажений'

                row3 = straight_params.rows[2]
                row3.cells[0].text = 'Наблюдатель'
                row3.cells[1].text = object_info.get('executor', '') if object_info else ''
                row3.cells[2].text = 'Дата'
                row3.cells[3].text = self.timestamp.split()[0] if ' ' in self.timestamp else self.timestamp
                row3.cells[4].text = ''
                row3.cells[5].text = ''

                format_table(
                    straight_params,
                    header_rows=0,
                    font_size=10,
                    column_widths=[3.0, 3.8, 3.0, 3.2, 3.0, 3.0]
                )

                doc.add_paragraph()

                # Заключение по прямолинейности
                max_location = (
                    f"\u043f\u043e\u044f\u0441 {straightness_summary['max_belt_num']}, "
                    f"\u0432\u044b\u0441\u043e\u0442\u0430 {straightness_summary['max_height_m']:.3f} \u043c"
                )
                if len(straightness_data) > 1 and straightness_summary['max_part_num'] is not None:
                    max_location = f"\u0447\u0430\u0441\u0442\u044c {straightness_summary['max_part_num']}, {max_location}"
                max_deviation = float(straightness_summary['max_deflection_mm'])

                straight_conclusion_heading = doc.add_paragraph('\u0417\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435:', style='Body Text')
                straight_conclusion_heading.runs[0].font.bold = True

                straight_conclusion = doc.add_paragraph(style='Body Text')
                straight_conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Straightness conclusion uses canonical profile points and their local tolerances
                tolerance = float(straightness_summary['max_tolerance_mm'])
                if straightness_summary['failed'] == 0:
                    straight_conclusion.add_run(
                        '\u0421\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430 \u043f\u043e\u044f\u0441\u043e\u0432 \u0431\u0430\u0448\u043d\u0438 \u043d\u0435 \u043f\u0440\u0435\u0432\u044b\u0448\u0430\u0435\u0442 '
                        '\u0434\u043e\u043f\u0443\u0441\u0442\u0438\u043c\u044b\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u044f. '
                        f'\u041f\u0440\u043e\u0432\u0435\u0440\u0435\u043d\u043e \u0442\u043e\u0447\u0435\u043a: {straightness_summary["total"]}. '
                        f'\u041c\u0430\u043a\u0441\u0438\u043c\u0430\u043b\u044c\u043d\u043e\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u0435 \u0441\u043e\u0441\u0442\u0430\u0432\u043b\u044f\u0435\u0442 {max_deviation:.1f} \u043c\u043c '
                        f'\u043f\u0440\u0438 \u0434\u043e\u043f\u0443\u0441\u043a\u0435 {tolerance:.1f} \u043c\u043c ({max_location}).'
                    )
                else:
                    run = straight_conclusion.add_run(
                        '\u0421\u0442\u0440\u0435\u043b\u0430 \u043f\u0440\u043e\u0433\u0438\u0431\u0430 \u043f\u043e\u044f\u0441\u043e\u0432 \u0431\u0430\u0448\u043d\u0438 \u043f\u0440\u0435\u0432\u044b\u0448\u0430\u0435\u0442 '
                        '\u0434\u043e\u043f\u0443\u0441\u0442\u0438\u043c\u044b\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u044f. '
                        f'\u041f\u0440\u0435\u0432\u044b\u0448\u0435\u043d\u0438\u0439: {straightness_summary["failed"]} \u0438\u0437 {straightness_summary["total"]} \u0442\u043e\u0447\u0435\u043a. '
                        f'\u041c\u0430\u043a\u0441\u0438\u043c\u0430\u043b\u044c\u043d\u043e\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u0435 \u0441\u043e\u0441\u0442\u0430\u0432\u043b\u044f\u0435\u0442 {max_deviation:.1f} \u043c\u043c '
                        f'\u043f\u0440\u0438 \u0434\u043e\u043f\u0443\u0441\u043a\u0435 {tolerance:.1f} \u043c\u043c ({max_location}).'
                    )
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True
                doc.add_paragraph()

                # Подписи
                straight_signature = doc.add_paragraph('Измерения выполнил              ', style='Body Text')
                straight_signature.add_run(object_info.get('executor', '____________') if object_info else '____________')
                straight_signature.add_run('        Вычисления проверил        ____________')

            # Сохраняем документ
            doc.save(output_path)

        except ImportError:
            raise ValueError("Для генерации DOCX отчетов требуется установить python-docx: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Ошибка генерации DOCX отчета: {e!s}")


